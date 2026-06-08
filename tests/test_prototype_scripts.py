from __future__ import annotations

import importlib.util
import io
import json
import sqlite3
import subprocess
import sys
from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document
try:
    from PIL import Image, ImageDraw
except ImportError:  # pragma: no cover - optional visual QA dependency
    Image = None
    ImageDraw = None


ROOT = Path(__file__).resolve().parents[1]
QA_SCRIPT = ROOT / "scripts" / "qa_document_artifacts.py"
DOCX_HELPER = ROOT / "tools" / "document_lib" / "keycentrix_docx.py"
GUIDE_DRAFT_SCRIPT = ROOT / "scripts" / "generate_guide_draft.py"
PROCESS_RECORDING_SCRIPT = ROOT / "scripts" / "process_recording.py"
BUILD_GUIDE_DOCX_SCRIPT = ROOT / "scripts" / "build_guide_docx.py"
COMPARE_TRANSCRIPTS_SCRIPT = ROOT / "scripts" / "compare_transcripts.py"
APP_SERVER_SCRIPT = ROOT / "scripts" / "app_server.py"


def load_module(path: Path, name: str):
    spec = importlib.util.spec_from_file_location(name, path)
    assert spec and spec.loader
    module = importlib.util.module_from_spec(spec)
    sys.modules[name] = module
    spec.loader.exec_module(module)
    return module


def write_docx(path: Path, paragraphs: list[str]) -> None:
    doc = Document()
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    doc.save(path)


def run_qa(path: Path) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        [sys.executable, str(QA_SCRIPT), "--json", str(path)],
        cwd=ROOT,
        check=False,
        text=True,
        capture_output=True,
    )


def test_artifact_qa_accepts_valid_guide_docx(tmp_path: Path) -> None:
    docx_path = tmp_path / "valid-guide.docx"
    write_docx(
        docx_path,
        [
            "keycentrix user guide",
            "Purpose",
            "Intended Audience",
            "Workflow Overview",
            "Step-by-Step Procedures",
            "Expected Results",
            "Troubleshooting",
            "Source Recording",
            "Screenshot reference: frame_0001.webp",
        ],
    )

    result = run_qa(docx_path)

    assert result.returncode == 0, result.stdout + result.stderr
    payload = json.loads(result.stdout)
    assert payload["passed"] is True
    assert payload["artifacts"][0]["missing_required_terms"] == []


def test_artifact_qa_rejects_forbidden_reference_terms(tmp_path: Path) -> None:
    docx_path = tmp_path / "leaky-guide.docx"
    write_docx(
        docx_path,
        [
            "keycentrix user guide",
            "Purpose",
            "Intended Audience",
            "Workflow Overview",
            "Step-by-Step Procedures",
            "Expected Results",
            "Troubleshooting",
            "Source Recording",
            "Screenshot reference: frame_0001.webp",
            "This draft accidentally mentions SmartReq.",
        ],
    )

    result = run_qa(docx_path)

    assert result.returncode == 1
    payload = json.loads(result.stdout)
    assert payload["passed"] is False
    assert "Reference-project terminology leaked." in payload["artifacts"][0]["forbidden_matches"][0]


def test_artifact_qa_strict_mode_rejects_prototype_placeholders(tmp_path: Path) -> None:
    docx_path = tmp_path / "placeholder-guide.docx"
    write_docx(
        docx_path,
        [
            "keycentrix user guide",
            "Purpose",
            "Intended Audience",
            "Workflow Overview",
            "Step-by-Step Procedures",
            "Expected Results",
            "Troubleshooting",
            "Source Recording",
            "Prototype narration segment 1",
        ],
    )

    result = subprocess.run(
        [sys.executable, str(QA_SCRIPT), "--json", "--strict", str(docx_path)],
        cwd=ROOT,
        check=False,
        text=True,
        capture_output=True,
    )

    assert result.returncode == 1
    payload = json.loads(result.stdout)
    assert "Prototype placeholder narration leaked." in payload["artifacts"][0]["forbidden_matches"][0]


def test_document_helper_can_build_minimal_docx_shell(tmp_path: Path) -> None:
    helper = load_module(DOCX_HELPER, "keycentrix_docx_test")
    output = tmp_path / "helper-guide.docx"

    doc = Document()
    helper.style_document(doc)
    helper.add_header_footer(doc, "Prototype Guide", "KCXDocumentor")
    helper.add_cover(
        doc,
        title="Prototype Guide",
        version="Version 0.1",
        status="Prototype",
        summary="Generated from local temporary test content.",
    )
    helper.add_section(
        doc,
        1,
        "Purpose",
        paragraphs=["Validate local DOCX creation without AI or external tools."],
    )
    doc.save(output)

    assert output.exists()
    with ZipFile(output) as package:
        assert "word/document.xml" in package.namelist()


def test_build_guide_docx_accepts_section_based_anthropic_shape(tmp_path: Path) -> None:
    input_path = tmp_path / "section-draft.json"
    output_path = tmp_path / "section-guide.docx"
    input_path.write_text(
        json.dumps(
            {
                "title": "SendKey User Guide",
                "meta": {
                    "draftStatus": "requires-human-review",
                    "sessionId": "sendkey-demo",
                    "targetApplication": "SendKey",
                },
                "introduction": {
                    "text": "SendKey provides communication methods for application users.",
                },
                "sections": [
                    {
                        "title": "Overview of Communication Methods",
                        "steps": [
                            {
                                "instruction": "Note that SendKey supports fax, SMS messaging, email, and voice messaging.",
                                "visibleUiText": ["SendKey"],
                                "reviewNotes": ["Confirm labels against OCR before publishing."],
                            }
                        ],
                    }
                ],
                "openReviewItems": [
                    {
                        "id": "review-001",
                        "severity": "blocking",
                        "description": "Screenshots require human approval.",
                        "resolution": "Approve candidate frames.",
                    }
                ],
                "model": {"provider": "anthropic"},
            }
        ),
        encoding="utf-8",
    )

    result = subprocess.run(
        [sys.executable, str(BUILD_GUIDE_DOCX_SCRIPT), str(input_path), "--output", str(output_path)],
        cwd=ROOT,
        check=False,
        text=True,
        capture_output=True,
    )

    assert result.returncode == 0, result.stdout + result.stderr
    assert output_path.exists()
    doc = Document(output_path)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "Overview of Communication Methods" in text
    assert "Note that SendKey supports fax" in text
    assert "Action:" not in text
    assert "Screenshots require human approval" not in text
    with ZipFile(output_path) as package:
        comments_xml = package.read("word/comments.xml").decode("utf-8")
    assert "Screenshots require human approval" in comments_xml


def test_build_guide_docx_converts_markdown_bold_to_word_runs(tmp_path: Path) -> None:
    input_path = tmp_path / "bold-draft.json"
    output_path = tmp_path / "bold-guide.docx"
    input_path.write_text(
        json.dumps(
            {
                "title": "Bold Guide",
                "summary": "Use **Blink RX** safely.",
                "workflow_overview": ["Open the **Interfaces** tab."],
                "steps": [
                    {
                        "title": "Configure Blink",
                        "instruction": "Set **Customer Type** to **Blink**.",
                        "expectedResult": "The **AR account** is selected.",
                    }
                ],
            }
        ),
        encoding="utf-8",
    )

    result = subprocess.run(
        [sys.executable, str(BUILD_GUIDE_DOCX_SCRIPT), str(input_path), "--output", str(output_path)],
        cwd=ROOT,
        check=False,
        text=True,
        capture_output=True,
    )

    assert result.returncode == 0, result.stdout + result.stderr
    doc = Document(output_path)
    paragraph_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "**" not in paragraph_text
    bold_run_text = {
        run.text
        for paragraph in doc.paragraphs
        for run in paragraph.runs
        if run.bold
    }
    assert {"Blink RX", "Interfaces", "Customer Type", "Blink", "AR account"}.issubset(bold_run_text)


def test_build_guide_docx_does_not_render_visible_screenshot_caption(tmp_path: Path) -> None:
    if Image is None:
        pytest.skip("Pillow is not installed")
    image_path = tmp_path / "application-frame.png"
    Image.new("RGB", (640, 360), "white").save(image_path)
    input_path = tmp_path / "caption-draft.json"
    output_path = tmp_path / "caption-guide.docx"
    input_path.write_text(
        json.dumps(
            {
                "title": "Screenshot Caption Guide",
                "steps": [
                    {
                        "title": "Open the request",
                        "instruction": "Open the request screen.",
                        "screenshot": str(image_path),
                        "caption": "Frame frame-0001 at 00:01:00.000",
                    }
                ],
            }
        ),
        encoding="utf-8",
    )

    result = subprocess.run(
        [sys.executable, str(BUILD_GUIDE_DOCX_SCRIPT), str(input_path), "--output", str(output_path)],
        cwd=ROOT,
        check=False,
        text=True,
        capture_output=True,
    )

    assert result.returncode == 0, result.stdout + result.stderr
    doc = Document(output_path)
    body_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "Frame frame-0001 at 00:01:00.000" not in body_text
    with ZipFile(output_path) as package:
        comments_xml = package.read("word/comments.xml").decode("utf-8")
    assert "Frame frame-0001 at 00:01:00.000" in comments_xml


def test_build_guide_docx_places_reviewer_concerns_in_comments_not_body(tmp_path: Path) -> None:
    input_path = tmp_path / "review-draft.json"
    output_path = tmp_path / "review-guide.docx"
    input_path.write_text(
        json.dumps(
            {
                "title": "Clean Review Guide",
                "summary": "Use this guide to complete the sample workflow.",
                "audience": ["Application users"],
                "workflow_overview": ["Complete the workflow in the order shown."],
                "expected_results": ["The workflow is completed successfully."],
                "steps": [
                    {
                        "title": "Save the record",
                        "instruction": "Select Save to finish updating the record.",
                        "expectedResult": "The record is saved.",
                        "reviewNotes": ["Confirm the screenshot does not include Teams controls."],
                        "visibleUiText": ["Save", "Teams"],
                        "confidence": {
                            "needsHumanReview": True,
                            "reasons": ["Transcript confidence is below publication threshold."],
                        },
                    }
                ],
                "openReviewItems": [
                    {
                        "id": "review-001",
                        "severity": "warning",
                        "description": "Screenshot selection requires approval.",
                    }
                ],
            }
        ),
        encoding="utf-8",
    )

    result = subprocess.run(
        [sys.executable, str(BUILD_GUIDE_DOCX_SCRIPT), str(input_path), "--output", str(output_path)],
        cwd=ROOT,
        check=False,
        text=True,
        capture_output=True,
    )

    assert result.returncode == 0, result.stdout + result.stderr
    doc = Document(output_path)
    body_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "Select Save to finish updating the record." in body_text
    assert "Confirm the screenshot does not include Teams controls." not in body_text
    assert "Transcript confidence is below publication threshold." not in body_text
    assert "Source UI evidence" not in body_text
    with ZipFile(output_path) as package:
        comments_xml = package.read("word/comments.xml").decode("utf-8")
    assert "Confirm the screenshot does not include Teams controls." in comments_xml
    assert "Transcript confidence is below publication threshold." in comments_xml


def test_build_guide_docx_cleans_redundant_step_headings_and_source_recording_metadata(tmp_path: Path) -> None:
    input_path = tmp_path / "blink-draft.json"
    output_path = tmp_path / "blink-guide.docx"
    input_path.write_text(
        json.dumps(
            {
                "title": "Blink Rx User Guide",
                "summary": "Use this guide to complete the Blink Rx workflow.",
                "sourceRecording": {
                    "sourceFile": "samples/raw/Blink Rx Training Part 2 120525.mp4",
                    "targetApplication": "Blink Rx",
                    "durationSeconds": 1242,
                },
                "sections": [
                    {
                        "title": "Submit a Refill Request",
                        "steps": [
                            {
                                "title": "Step 5 — Submit a Refill Request: Step 14",
                                "instruction": "Select Submit to send the refill request.",
                                "expectedResult": "The refill request is submitted.",
                            }
                        ],
                    }
                ],
            }
        ),
        encoding="utf-8",
    )

    result = subprocess.run(
        [sys.executable, str(BUILD_GUIDE_DOCX_SCRIPT), str(input_path), "--output", str(output_path)],
        cwd=ROOT,
        check=False,
        text=True,
        capture_output=True,
    )

    assert result.returncode == 0, result.stdout + result.stderr
    doc = Document(output_path)
    body_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    assert "1. Submit a Refill Request" in body_text
    assert "Submit a Refill Request: Step 14" not in body_text
    assert "Application\nBlink Rx" in table_text
    assert "Application\nNot specified" not in table_text


def test_build_guide_docx_uses_document_metadata_from_anthropic_draft(tmp_path: Path) -> None:
    input_path = tmp_path / "document-metadata-draft.json"
    output_path = tmp_path / "document-metadata-guide.docx"
    input_path.write_text(
        json.dumps(
            {
                "schemaVersion": 1,
                "sessionId": "blink-rx-part-2-sidecar",
                "document": {
                    "title": "Blink Rx Workflow Guide",
                    "targetApplication": "Blink Rx",
                    "durationSeconds": 1241.601,
                    "sourceRecording": "Blink Rx Training Part 2 120525.mp4",
                },
                "sections": [
                    {
                        "title": "Submit a Refill Request",
                        "steps": [
                            {
                                "title": "Initiate the refill request",
                                "instruction": "Set the refill initiation flag to true.",
                            }
                        ],
                    }
                ],
            }
        ),
        encoding="utf-8",
    )

    result = subprocess.run(
        [sys.executable, str(BUILD_GUIDE_DOCX_SCRIPT), str(input_path), "--output", str(output_path)],
        cwd=ROOT,
        check=False,
        text=True,
        capture_output=True,
    )

    assert result.returncode == 0, result.stdout + result.stderr
    doc = Document(output_path)
    table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    assert "Application\nBlink Rx" in table_text
    assert "Recording Duration\n00:20:42" in table_text
    assert "Recording Duration\n1241.601" not in table_text
    assert "Source Recording\nBlink Rx Training Part 2 120525.mp4" in table_text


def test_build_guide_docx_replaces_internal_purpose_and_prerequisite_language(tmp_path: Path) -> None:
    input_path = tmp_path / "customer-safe-draft.json"
    output_path = tmp_path / "customer-safe-guide.docx"
    input_path.write_text(
        json.dumps(
            {
                "title": "Blink Rx Workflow Guide",
                "summary": "This guide was generated from a local procedure trace.",
                "document": {
                    "description": "This guide covers refill requests and pharmacy profile plan template configuration in Blink Rx.",
                    "targetApplication": "Blink Rx",
                },
                "prerequisites": [
                    "Confirm access to the target application and review the source recording context before publishing."
                ],
                "steps": [
                    {
                        "title": "Open the refill workflow",
                        "instruction": "Open the refill request workflow.",
                    }
                ],
            }
        ),
        encoding="utf-8",
    )

    result = subprocess.run(
        [sys.executable, str(BUILD_GUIDE_DOCX_SCRIPT), str(input_path), "--output", str(output_path)],
        cwd=ROOT,
        check=False,
        text=True,
        capture_output=True,
    )

    assert result.returncode == 0, result.stdout + result.stderr
    doc = Document(output_path)
    body_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "This guide covers refill requests and pharmacy profile plan template configuration in Blink Rx." in body_text
    assert "generated from a local procedure trace" not in body_text
    assert "before publishing" not in body_text
    assert "application reviewer" not in body_text
    assert "You have access to Blink Rx and the permissions needed to complete this workflow." in body_text


def test_build_guide_docx_moves_candidate_screenshot_language_to_comments(tmp_path: Path) -> None:
    image_path = tmp_path / "workflow.png"
    if Image is not None:
        Image.new("RGB", (320, 180), color=(245, 245, 245)).save(image_path)
    else:
        image_path.write_bytes(b"not-a-real-image")
    input_path = tmp_path / "candidate-language-draft.json"
    output_path = tmp_path / "candidate-language-guide.docx"
    input_path.write_text(
        json.dumps(
            {
                "title": "Workflow Review Guide",
                "status": "Prototype",
                "summary": "This prototype guide covers the candidate screenshot workflow.",
                "steps": [
                    {
                        "title": "Candidate screenshot must be selected",
                        "instruction": "Candidate screenshot at 00:01:00 shows the settings page.",
                        "expectedResult": "No screenshot was available for this procedure step.",
                        "screenshot": str(image_path),
                        "caption": "Candidate screenshot at 00:01:00",
                    }
                ],
            }
        ),
        encoding="utf-8",
    )

    result = subprocess.run(
        [sys.executable, str(BUILD_GUIDE_DOCX_SCRIPT), str(input_path), "--output", str(output_path)],
        cwd=ROOT,
        check=False,
        text=True,
        capture_output=True,
    )

    assert result.returncode == 0, result.stdout + result.stderr
    doc = Document(output_path)
    body_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "Prototype" not in body_text
    assert "prototype" not in body_text
    assert "Candidate screenshot" not in body_text
    assert "candidate screenshot" not in body_text
    assert "No screenshot was available" not in body_text
    assert "Final Revision for Review" in body_text
    with ZipFile(output_path) as package:
        comments_xml = package.read("word/comments.xml").decode("utf-8")
    assert "Candidate screenshot at 00:01:00" in comments_xml
    assert "No screenshot was available for this procedure step." in comments_xml


def test_build_guide_docx_infers_contextual_audience(tmp_path: Path) -> None:
    input_path = tmp_path / "audience-draft.json"
    output_path = tmp_path / "audience-guide.docx"
    input_path.write_text(
        json.dumps(
            {
                "title": "Blink Rx Training Guide",
                "document": {"targetApplication": "Blink Rx"},
                "audience": ["Application users"],
                "summary": "This guide covers refill and pharmacy profile workflows in Blink Rx.",
                "steps": [{"title": "Open the profile", "instruction": "Open the pharmacy profile."}],
            }
        ),
        encoding="utf-8",
    )

    result = subprocess.run(
        [sys.executable, str(BUILD_GUIDE_DOCX_SCRIPT), str(input_path), "--output", str(output_path)],
        cwd=ROOT,
        check=False,
        text=True,
        capture_output=True,
    )

    assert result.returncode == 0, result.stdout + result.stderr
    body_text = "\n".join(paragraph.text for paragraph in Document(output_path).paragraphs)
    assert "Application users" not in body_text
    assert "Trainers and team members learning the workflow" in body_text
    assert "Pharmacy workflow team members who complete or support the documented process" in body_text


def test_build_guide_docx_omits_generation_metadata_from_delivered_docx(tmp_path: Path) -> None:
    input_path = tmp_path / "usage-draft.json"
    output_path = tmp_path / "usage-guide.docx"
    input_path.write_text(
        json.dumps(
            {
                "generatedAt": "2026-06-04T16:34:09Z",
                "model": "claude-sonnet-4-6",
                "usage": {
                    "input_tokens": 12847,
                    "output_tokens": 3201,
                    "totalTokens": 16048,
                    "estimatedCostUSD": 0.086,
                },
                "guideDraft": {
                    "title": "Blink Rx Workflow Guide",
                    "summary": "This guide covers refill requests in Blink Rx.",
                    "sourceRecording": {"targetApplication": "Blink Rx"},
                    "steps": [
                        {
                            "title": "Open the refill workflow",
                            "instruction": "Open the refill request workflow.",
                        }
                    ],
                },
            }
        ),
        encoding="utf-8",
    )

    result = subprocess.run(
        [sys.executable, str(BUILD_GUIDE_DOCX_SCRIPT), str(input_path), "--output", str(output_path)],
        cwd=ROOT,
        check=False,
        text=True,
        capture_output=True,
    )

    assert result.returncode == 0, result.stdout + result.stderr
    doc = Document(output_path)
    headings = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    assert "Appendix: Generation Metadata" not in headings
    assert "Input Tokens" not in table_text
    assert "Output Tokens" not in table_text
    assert "Total Tokens" not in table_text
    assert "Estimated Cost" not in table_text


def test_build_guide_docx_adds_reviewer_comments_for_ambiguous_terms_without_body_pollution(tmp_path: Path) -> None:
    input_path = tmp_path / "ambiguous-terms-draft.json"
    output_path = tmp_path / "ambiguous-terms-guide.docx"
    input_path.write_text(
        json.dumps(
            {
                "title": "Blink Rx Workflow Guide",
                "summary": "This guide covers the Blink Rx refill workflow.",
                "sourceRecording": {"targetApplication": "Blink Rx"},
                "steps": [
                    {
                        "title": "Confirm required checkbox selections",
                        "instruction": "Select all required checkboxes before choosing PV1 and PDR.",
                        "expectedResult": "The refill request is ready for review.",
                    }
                ],
            }
        ),
        encoding="utf-8",
    )

    result = subprocess.run(
        [sys.executable, str(BUILD_GUIDE_DOCX_SCRIPT), str(input_path), "--output", str(output_path)],
        cwd=ROOT,
        check=False,
        text=True,
        capture_output=True,
    )

    assert result.returncode == 0, result.stdout + result.stderr
    doc = Document(output_path)
    body_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "Select all required checkboxes before choosing PV1 and PDR." in body_text
    assert "Verify that 'PV1'" not in body_text
    assert "Verify the exact required checkbox labels" not in body_text
    with ZipFile(output_path) as package:
        comments_xml = package.read("word/comments.xml").decode("utf-8")
    assert "Verify that 'PV1' is the correct UI term" in comments_xml
    assert "Verify that 'PDR' is the correct UI term" in comments_xml
    assert "Verify the exact required checkbox labels" in comments_xml


def test_build_guide_docx_routes_dark_meeting_frame_to_reviewer_comment(tmp_path: Path) -> None:
    if Image is None:
        pytest.skip("Pillow is not installed")
    image_path = tmp_path / "teams-card.png"
    Image.new("RGB", (320, 180), color=(0, 0, 0)).save(image_path)
    input_path = tmp_path / "dark-frame-draft.json"
    output_path = tmp_path / "dark-frame-guide.docx"
    input_path.write_text(
        json.dumps(
            {
                "title": "Dark Frame Guide",
                "summary": "Use this guide to complete the workflow.",
                "steps": [
                    {
                        "title": "Understand the template setting",
                        "instruction": "Review the plan template behavior.",
                        "screenshot": str(image_path),
                    }
                ],
            }
        ),
        encoding="utf-8",
    )

    result = subprocess.run(
        [sys.executable, str(BUILD_GUIDE_DOCX_SCRIPT), str(input_path), "--output", str(output_path)],
        cwd=ROOT,
        check=False,
        text=True,
        capture_output=True,
    )

    assert result.returncode == 0, result.stdout + result.stderr
    doc = Document(output_path)
    body_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "Candidate screenshot" not in body_text
    assert "No screenshot was available" not in body_text
    assert "non-application Teams or meeting frame" not in body_text
    with ZipFile(output_path) as package:
        comments_xml = package.read("word/comments.xml").decode("utf-8")
        document_xml = package.read("word/document.xml").decode("utf-8")
    assert "non-application Teams or meeting frame" in comments_xml
    assert "teams-card.png" not in document_xml


def test_build_guide_docx_routes_centered_meeting_overlay_to_reviewer_comment(tmp_path: Path) -> None:
    if Image is None or ImageDraw is None:
        pytest.skip("Pillow is not installed")
    module = load_module(BUILD_GUIDE_DOCX_SCRIPT, "build_guide_docx_centered_overlay")
    image_path = tmp_path / "meeting-overlay.png"
    image = Image.new("RGB", (900, 520), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((270, 170, 630, 350), fill=(20, 20, 24))
    image.save(image_path)

    assert module.is_likely_non_application_screenshot(image_path) is True


def test_build_guide_docx_preserves_detailed_review_guidance_in_comments(tmp_path: Path) -> None:
    input_path = tmp_path / "review-guidance-draft.json"
    output_path = tmp_path / "review-guidance-guide.docx"
    input_path.write_text(
        json.dumps(
            {
                "title": "Guidance Comment Guide",
                "summary": "Use this guide to complete the workflow.",
                "reviewGuidance": [
                    {
                        "id": "review-021",
                        "severity": "warning",
                        "description": "Verify segment seg-0021 because the narrator describes the pharmacy profile template quickly.",
                        "resolution": "Review the live application if the selected frame does not show the template fields.",
                    }
                ],
                "steps": [
                    {
                        "title": "Review the pharmacy profile template",
                        "instruction": "Review the pharmacy profile template fields.",
                        "expectedResult": "The profile template is ready for review.",
                        "reviewerGuidance": [
                            "Segment seg-0021 has low transcript confidence; confirm plan-template terminology before publishing."
                        ],
                    }
                ],
            }
        ),
        encoding="utf-8",
    )

    result = subprocess.run(
        [sys.executable, str(BUILD_GUIDE_DOCX_SCRIPT), str(input_path), "--output", str(output_path)],
        cwd=ROOT,
        check=False,
        text=True,
        capture_output=True,
    )

    assert result.returncode == 0, result.stdout + result.stderr
    doc = Document(output_path)
    body_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "low transcript confidence" not in body_text
    assert "review-021" not in body_text
    with ZipFile(output_path) as package:
        comments_xml = package.read("word/comments.xml").decode("utf-8")
    assert "Segment seg-0021 has low transcript confidence" in comments_xml
    assert "review-021" in comments_xml
    assert "pharmacy profile template" in comments_xml


def test_docx_builder_resolves_processed_session_frame_assets(tmp_path: Path) -> None:
    module = load_module(BUILD_GUIDE_DOCX_SCRIPT, "build_guide_docx_asset_roots")
    session_dir = tmp_path / "samples" / "processed" / "demo-session"
    image_path = session_dir / "frames" / "candidates" / "frame-0001.png"
    image_path.parent.mkdir(parents=True)
    image_path.write_bytes(b"not-a-real-png")
    draft_path = tmp_path / "artifacts" / "generated" / "demo-session" / "guide_draft.json"
    draft_path.parent.mkdir(parents=True)

    resolved = module.resolve_asset_path("frames/candidates/frame-0001.png", draft_path, [session_dir])

    assert resolved == image_path.resolve()


def test_artifact_qa_reports_reviewer_comments_and_body_cleanliness(tmp_path: Path) -> None:
    input_path = tmp_path / "commented-draft.json"
    output_path = tmp_path / "commented-guide.docx"
    input_path.write_text(
        json.dumps(
            {
                "title": "Commented Guide",
                "summary": "Use this guide to complete the workflow.",
                "audience": ["Application users"],
                "workflow_overview": ["Complete the documented task."],
                "expected_results": ["The task is complete."],
                "steps": [
                    {
                        "title": "Review and save",
                        "instruction": "Review the record and select Save.",
                        "expectedResult": "The record is saved.",
                        "reviewNotes": ["Visible UI text pending local OCR should be resolved before release."],
                    }
                ],
            }
        ),
        encoding="utf-8",
    )
    build = subprocess.run(
        [sys.executable, str(BUILD_GUIDE_DOCX_SCRIPT), str(input_path), "--output", str(output_path)],
        cwd=ROOT,
        check=False,
        text=True,
        capture_output=True,
    )
    assert build.returncode == 0, build.stdout + build.stderr

    result = subprocess.run(
        [sys.executable, str(QA_SCRIPT), "--json", "--strict", str(output_path)],
        cwd=ROOT,
        check=False,
        text=True,
        capture_output=True,
    )

    assert result.returncode == 0, result.stdout + result.stderr
    payload = json.loads(result.stdout)
    artifact = payload["artifacts"][0]
    assert artifact["body_clean"] is True
    assert artifact["reviewer_comment_count"] >= 1
    assert any("reviewer comments" in warning.lower() for warning in artifact["warnings"])


def test_build_guide_docx_sanitizes_prompt_language_in_reviewer_comments(tmp_path: Path) -> None:
    input_path = tmp_path / "leaky-comment-draft.json"
    output_path = tmp_path / "leaky-comment-guide.docx"
    input_path.write_text(
        json.dumps(
            {
                "title": "Leaky Comment Guide",
                "summary": "Use this guide to complete the workflow.",
                "audience": ["Application users"],
                "workflow_overview": ["Complete the documented task."],
                "expected_results": ["The task is complete."],
                "steps": [
                    {
                        "title": "Review and save",
                        "instruction": "Review the record and select Save.",
                        "expectedResult": "The record is saved.",
                        "reviewNotes": ["The system prompt should decide this action."],
                    }
                ],
            }
        ),
        encoding="utf-8",
    )
    build = subprocess.run(
        [sys.executable, str(BUILD_GUIDE_DOCX_SCRIPT), str(input_path), "--output", str(output_path)],
        cwd=ROOT,
        check=False,
        text=True,
        capture_output=True,
    )
    assert build.returncode == 0, build.stdout + build.stderr

    result = subprocess.run(
        [sys.executable, str(QA_SCRIPT), "--json", str(output_path)],
        cwd=ROOT,
        check=False,
        text=True,
        capture_output=True,
    )

    assert result.returncode == 0, result.stdout + result.stderr
    with ZipFile(output_path) as package:
        comments_xml = package.read("word/comments.xml").decode("utf-8")
    assert "system prompt" not in comments_xml
    assert "application message" in comments_xml


def test_artifact_qa_rejects_prompt_or_reasoning_leaks_in_comments(tmp_path: Path) -> None:
    docx_path = tmp_path / "raw-leaky-comment.docx"
    doc = Document()
    for paragraph in [
        "keycentrix user guide",
        "Purpose",
        "Intended Audience",
        "Workflow Overview",
        "Step-by-Step Procedures",
        "Expected Results",
        "Troubleshooting",
        "Source Recording",
    ]:
        doc.add_paragraph(paragraph)
    anchor = doc.add_paragraph("Review and save the record.")
    doc.add_comment(anchor.runs, text="The system prompt should decide this action.", author="KCXDocumentor Reviewer")
    doc.save(docx_path)

    result = subprocess.run(
        [sys.executable, str(QA_SCRIPT), "--json", str(docx_path)],
        cwd=ROOT,
        check=False,
        text=True,
        capture_output=True,
    )

    assert result.returncode == 1
    payload = json.loads(result.stdout)
    assert "Internal prompt terminology leaked." in payload["artifacts"][0]["forbidden_matches"][0]


def test_compact_procedure_trace_contract_for_one_hour_recordings(tmp_path: Path) -> None:
    trace_path = tmp_path / "procedure_trace.json"
    trace = {
        "session": {
            "app_name": "Sample Enterprise App",
            "duration_sec": 3600,
            "recorded_at": "2026-06-04T10:00:00-05:00",
        },
        "steps": [
            {
                "start": "00:02:10.000",
                "end": "00:03:05.250",
                "speaker_text": "Open the customer record and click Save.",
                "visible_ui_text": ["Customer", "Save"],
                "action_hints": ["open_record", "mouse_click"],
                "candidate_images": ["frames/frame_000130.webp"],
                "confidence": {
                    "transcript": 0.91,
                    "ocr": 0.84,
                    "frameSelection": 0.88,
                    "overall": 0.88,
                    "needsHumanReview": False,
                    "reasons": [],
                },
            }
        ],
    }
    trace_path.write_text(json.dumps(trace), encoding="utf-8")

    loaded = json.loads(trace_path.read_text(encoding="utf-8"))

    assert loaded["session"]["duration_sec"] >= 3600
    assert len(json.dumps(loaded)) < 1200
    assert loaded["steps"][0]["speaker_text"]
    assert loaded["steps"][0]["confidence"]["overall"] >= 0.75
    assert loaded["steps"][0]["visible_ui_text"] == ["Customer", "Save"]
    assert loaded["steps"][0]["candidate_images"][0].endswith(".webp")


def test_guide_draft_generator_requires_remote_proxy_url(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    module = load_module(GUIDE_DRAFT_SCRIPT, "generate_guide_draft_anthropic_required")
    trace = {
        "schemaVersion": 1,
        "recording": {
            "sourceFile": "samples/raw/example.mp4",
            "durationSeconds": 3600,
            "targetApplication": "Enterprise Rx",
            "captureMode": "imported-recording",
        },
        "segments": [
            {
                "id": "seg-0001",
                "start": "00:00:00.000",
                "end": "00:01:00.000",
                "speakerText": "I click Save to finish the record.",
                "visibleUiText": ["Save"],
                "actionHints": ["click", "save"],
                "candidateImages": [{"path": "", "timestamp": "00:00:30.000", "reason": "candidate", "reviewStatus": "pending"}],
                "confidence": {
                    "overall": 0.42,
                    "needsHumanReview": True,
                    "reasons": ["Transcript confidence is below publication threshold."],
                },
            }
        ],
    }
    monkeypatch.delenv("KCXDOC_REMOTE_API_BASE_URL", raising=False)

    with pytest.raises(SystemExit, match="KCXDOC_REMOTE_API_BASE_URL is required"):
        module.generate_with_anthropic(
            trace,
            type(
                "Args",
                (),
                {
                    "model": "claude-sonnet-4-6",
                    "max_tokens": 100,
                    "temperature": 0.2,
                    "prompt_version": "guide-draft-v1",
                },
            )(),
        )


def test_prepare_trace_for_anthropic_prunes_rejected_frames_and_preserves_review_guidance() -> None:
    module = load_module(GUIDE_DRAFT_SCRIPT, "generate_guide_draft_prepare_trace")
    trace = {
        "schemaVersion": 1,
        "reviewGuidance": [{"type": "existing", "message": "Confirm terminology."}],
        "segments": [
            {
                "id": "seg-0001",
                "candidateImages": [
                    {
                        "frameId": "frame-approved",
                        "path": "frames/frame-approved.png",
                        "timestamp": "00:01:00.000",
                        "timestampSeconds": 60.0,
                        "created": True,
                        "reviewStatus": "approved",
                    },
                    {
                        "frameId": "frame-rejected",
                        "path": "frames/frame-rejected.png",
                        "timestamp": "00:01:30.000",
                        "timestampSeconds": 90.0,
                        "created": True,
                        "reviewStatus": "rejected",
                        "reviewNote": "Do not use this Teams title card.",
                    },
                ],
            }
        ],
    }

    prepared = module.prepare_trace_for_anthropic(trace)
    candidate_ids = [
        image["frameId"]
        for segment in prepared["segments"]
        for image in segment["candidateImages"]
    ]

    assert candidate_ids == ["frame-approved"]
    assert trace["segments"][0]["candidateImages"][1]["frameId"] == "frame-rejected"
    assert prepared["excludedFrames"][0]["frameId"] == "frame-rejected"
    assert prepared["excludedFrames"][0]["sourceSegmentId"] == "seg-0001"
    assert "Do not use this Teams title card." in json.dumps(prepared["reviewGuidance"])
    assert "Confirm terminology." in json.dumps(prepared["reviewGuidance"])


def test_prepare_trace_for_anthropic_prunes_system_rejected_frames() -> None:
    module = load_module(GUIDE_DRAFT_SCRIPT, "generate_guide_draft_system_rejected")
    trace = {
        "segments": [
            {
                "id": "seg-0001",
                "candidateImages": [
                    {
                        "frameId": "frame-good",
                        "path": "frames/frame-good.png",
                        "timestamp": "00:01:00.000",
                        "timestampSeconds": 60.0,
                        "created": True,
                        "reviewStatus": "pending",
                        "recommendationGroup": "recommended",
                        "frameEvidenceScore": 0.82,
                    },
                    {
                        "frameId": "frame-system-rejected",
                        "path": "frames/frame-system-rejected.png",
                        "timestamp": "00:01:30.000",
                        "timestampSeconds": 90.0,
                        "created": True,
                        "reviewStatus": "pending",
                        "recommendationGroup": "system-rejected",
                        "recommendationReason": "Likely Teams title card.",
                    },
                ],
            }
        ],
    }

    prepared = module.prepare_trace_for_anthropic(trace)

    assert [image["frameId"] for image in prepared["segments"][0]["candidateImages"]] == ["frame-good"]
    assert prepared["excludedFrames"][0]["frameId"] == "frame-system-rejected"
    assert prepared["excludedFrames"][0]["exclusionReason"] == "Likely Teams title card."


def test_prepare_trace_for_anthropic_compacts_ocr_payload() -> None:
    module = load_module(GUIDE_DRAFT_SCRIPT, "generate_guide_draft_compact_ocr")
    long_ocr = " | ".join(
        [
            "Blink Mock UI Refill Request Patient Profile Submit",
            "12345",
            "----",
            "Transfer Back Transfer In Claims Dispense Create Request",
        ]
        * 30
    )
    trace = {
        "segments": [
            {
                "id": "seg-0001",
                "visibleUiText": ["1", "-", "Blink Mock UI", "Submit", "Submit", "12345"],
                "candidateImages": [
                    {
                        "frameId": "frame-1",
                        "path": "frames/frame-1.png",
                        "timestamp": "00:01:00.000",
                        "timestampSeconds": 60.0,
                        "created": True,
                        "reviewStatus": "pending",
                        "frameEvidenceScore": 0.8,
                        "ocrConfidence": 0.7,
                        "recommendationGroup": "recommended",
                        "selectionReasons": ["OCR indicates the target application."],
                        "contentType": "application",
                        "ocrText": long_ocr,
                        "reason": "A very long reason. " * 40,
                    },
                    {
                        "frameId": "frame-2",
                        "path": "frames/frame-2.png",
                        "timestamp": "00:01:20.000",
                        "timestampSeconds": 80.0,
                        "created": True,
                        "reviewStatus": "pending",
                        "frameEvidenceScore": 0.4,
                        "ocrText": long_ocr,
                    },
                    {
                        "frameId": "frame-3",
                        "path": "frames/frame-3.png",
                        "timestamp": "00:01:40.000",
                        "timestampSeconds": 100.0,
                        "created": True,
                        "reviewStatus": "pending",
                        "frameEvidenceScore": 0.3,
                        "ocrText": long_ocr,
                    },
                    {
                        "frameId": "frame-4",
                        "path": "frames/frame-4.png",
                        "timestamp": "00:02:00.000",
                        "timestampSeconds": 120.0,
                        "created": True,
                        "reviewStatus": "pending",
                        "frameEvidenceScore": 0.2,
                        "ocrText": long_ocr,
                    },
                ],
            }
        ]
    }

    prepared = module.prepare_trace_for_anthropic(trace)
    segment = prepared["segments"][0]

    assert segment["visibleUiText"] == ["Blink Mock UI", "Submit"]
    assert len(segment["candidateImages"]) == 3
    assert segment["candidateImages"][0]["frameId"] == "frame-1"
    assert segment["candidateImages"][0]["recommendationGroup"] == "recommended"
    assert segment["candidateImages"][0]["contentType"] == "application"
    assert len(segment["candidateImages"][0]["ocrText"]) <= 360
    assert len(segment["candidateImages"][0]["reason"]) <= 220


def test_generation_quality_rules_block_zero_transcript_without_segment_steps() -> None:
    module = load_module(GUIDE_DRAFT_SCRIPT, "generate_guide_draft_zero_transcript_rules")
    trace = {
        "sessionId": "zero-transcript",
        "recording": {
            "sourceFile": "samples/raw/KCX Bootcamp-Retail Specialty Workflow.mp4",
            "durationSeconds": 2068,
            "captureMode": "imported-recording",
        },
        "segments": [
            {"id": f"seg-{index:04d}", "speakerText": "", "confidence": {"transcript": 0.0}, "candidateImages": []}
            for index in range(27)
        ],
    }
    draft = {
        "schemaVersion": 1,
        "title": "Retail Specialty Workflow",
        "steps": [{"title": "Bad model segment step"} for _ in range(27)],
    }

    normalized = module.enforce_generation_quality_rules(draft, trace)

    assert normalized["overallStatus"] == "BLOCKED — No transcript available."
    assert len(normalized["steps"]) == 4
    assert len(normalized["steps"]) < len(trace["segments"])
    assert {step["title"] for step in normalized["steps"]} == {"PLACEHOLDER — requires transcript"}
    assert {step["instruction"] for step in normalized["steps"]} == {"PLACEHOLDER — requires transcript"}
    review = normalized["openReviewItems"][0]
    assert review["id"] == "review-001"
    assert review["severity"] == "critical"
    assert review["totalSegmentCount"] == 27
    assert review["recordingDuration"] == "34 minutes 28 seconds"
    assert "No usable transcript was extracted" in review["description"]
    assert normalized["sourceRecording"] == {
        "fileName": "KCX Bootcamp-Retail Specialty Workflow.mp4",
        "sourceFile": "samples/raw/KCX Bootcamp-Retail Specialty Workflow.mp4",
        "duration": "34 minutes 28 seconds",
        "durationSeconds": 2068.0,
        "captureMode": "imported-recording",
    }


def test_generation_quality_rules_consolidate_systemic_screenshot_failure_and_cadence() -> None:
    module = load_module(GUIDE_DRAFT_SCRIPT, "generate_guide_draft_screenshot_consolidation")
    trace = {
        "recording": {
            "sourceFile": "samples/raw/demo.mp4",
            "durationSeconds": 3600,
            "captureMode": "imported-recording",
        },
        "segments": [
            {
                "id": "seg-0001",
                "speakerText": "Open the request.",
                "confidence": {"transcript": 0.9},
                "candidateImages": [
                    {
                        "frameId": "frame-0001",
                        "timestampSeconds": 60,
                        "reviewStatus": "pending",
                        "confidence": 0.9,
                        "ocrText": "Open Request",
                    }
                ],
            },
            {
                "id": "seg-0002",
                "speakerText": "Submit it.",
                "confidence": {"transcript": 0.9},
                "candidateImages": [
                    {
                        "frameId": "frame-0002",
                        "timestampSeconds": 90,
                        "reviewStatus": "pending",
                        "confidence": 0.9,
                        "ocrText": "Submit",
                    }
                ],
            },
        ],
    }
    draft = {
        "schemaVersion": 1,
        "title": "Demo Guide",
        "sourceRecording": {"sourceFile": "demo.mp4", "durationSeconds": 3600},
        "steps": [
            {"title": "Open request", "instruction": "Open the request.", "selectedScreenshot": {"timestampSeconds": 60}},
            {"title": "Submit request", "instruction": "Submit the request.", "selectedScreenshot": {"timestampSeconds": 90}},
        ],
    }

    normalized = module.enforce_generation_quality_rules(draft, trace)

    review_by_id = {item["id"]: item for item in normalized["openReviewItems"]}
    assert review_by_id["review-001"]["recordingDuration"] == "1 hour"
    assert review_by_id["review-002"]["severity"] == "critical"
    assert review_by_id["review-002"]["description"] == "All candidate screenshots are pending review."
    assert "cadence-based" in review_by_id["review-003"]["description"]
    for step in normalized["steps"]:
        assert step["screenshotDecision"] == {
            "needsHumanReview": True,
            "reviewNote": "See review-002.",
            "screenshotRef": None,
        }


def test_anthropic_payload_excludes_rejected_frame_candidates_but_keeps_notes(monkeypatch: pytest.MonkeyPatch) -> None:
    module = load_module(GUIDE_DRAFT_SCRIPT, "generate_guide_draft_payload_contract")
    captured: dict[str, object] = {}

    class FakeResponse:
        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, traceback) -> None:
            return None

        def read(self) -> bytes:
            return json.dumps(
                {
                    "anthropicResult": {
                        "content": [
                            {
                                "type": "text",
                                "text": json.dumps({"title": "Guide", "steps": []}),
                            }
                        ],
                        "usage": {"input_tokens": 12847, "output_tokens": 3201},
                    },
                    "generationReport": {
                        "generationRunId": "run-123",
                        "generatedAt": "2026-06-04T16:34:09Z",
                    },
                }
            ).encode("utf-8")

    def fake_urlopen(req, timeout):
        payload = json.loads(req.data.decode("utf-8"))
        user_prompt = json.loads(payload["anthropic"]["messages"][0]["content"])
        captured["procedureTrace"] = user_prompt["procedureTrace"]
        return FakeResponse()

    trace = {
        "sessionId": "demo-session",
        "segments": [
            {
                "id": "seg-0001",
                "candidateImages": [
                    {
                        "frameId": "keep-me",
                        "path": "frames/keep-me.png",
                        "created": True,
                        "reviewStatus": "pending",
                    },
                    {
                        "frameId": "reject-me",
                        "path": "frames/reject-me.png",
                        "created": True,
                        "reviewStatus": "rejected",
                        "reviewNote": "Rejected because it shows meeting controls.",
                    },
                ],
            }
        ],
    }
    monkeypatch.setenv("KCXDOC_REMOTE_API_BASE_URL", "https://kcxdocumentor-ai-dev.azurewebsites.net")
    monkeypatch.setenv("KCXDOC_REMOTE_API_BEARER_TOKEN", "token-123")
    monkeypatch.setattr(module.request, "urlopen", fake_urlopen)

    draft = module.generate_with_anthropic(
        trace,
        type(
            "Args",
            (),
            {
                "model": "claude-sonnet-4-6",
                "max_tokens": 100,
                "temperature": 0.2,
                "prompt_version": "guide-draft-v1",
            },
        )(),
    )

    procedure_trace = captured["procedureTrace"]
    candidate_ids = [
        image["frameId"]
        for segment in procedure_trace["segments"]
        for image in segment["candidateImages"]
    ]
    assert candidate_ids == ["keep-me"]
    assert procedure_trace["excludedFrames"][0]["frameId"] == "reject-me"
    assert "Rejected because it shows meeting controls." in json.dumps(procedure_trace["reviewGuidance"])
    assert draft["generatedAt"].endswith("Z")
    assert draft["generationRunId"] == "run-123"
    assert draft["model"]["model"] == "claude-sonnet-4-6"
    assert draft["usage"] == {
        "inputTokens": 12847,
        "outputTokens": 3201,
        "totalTokens": 16048,
        "estimatedCostUSD": 0.086556,
    }


def test_generation_report_is_written_next_to_anthropic_draft(tmp_path: Path) -> None:
    module = load_module(GUIDE_DRAFT_SCRIPT, "generate_guide_draft_generation_report")
    draft_path = tmp_path / "session" / "guide_draft.anthropic.json"
    module.USAGE_DB_PATH = tmp_path / "usage" / "generation_usage.sqlite3"
    draft_path.parent.mkdir(parents=True)
    draft = {
        "generatedAt": "2026-06-04T16:34:09Z",
        "sessionId": "demo-session",
        "document": {"title": "Blink Rx Workflow Guide"},
        "model": {
            "provider": "anthropic",
            "model": "claude-sonnet-4-6",
            "promptVersion": "guide-draft-v1",
        },
        "usage": {
            "inputTokens": 12847,
            "outputTokens": 3201,
            "totalTokens": 16048,
            "estimatedCostUSD": 0.086556,
        },
    }

    report_path = module.write_generation_report(draft_path, draft)

    assert report_path == draft_path.parent / "generation_report.json"
    report = json.loads(report_path.read_text(encoding="utf-8"))
    expected = {
        "schemaVersion": 1,
        "status": "succeeded",
        "generatedAt": "2026-06-04T16:34:09Z",
        "sessionId": "demo-session",
        "title": "Blink Rx Workflow Guide",
        "model": "claude-sonnet-4-6",
        "provider": "anthropic",
        "promptVersion": "guide-draft-v1",
        "usage": {
            "inputTokens": 12847,
            "outputTokens": 3201,
            "totalTokens": 16048,
            "estimatedCostUSD": 0.086556,
        },
    }
    expected["generationRunId"] = module.generation_run_id(expected)
    assert report == expected
    assert not module.USAGE_DB_PATH.exists()


def test_anthropic_invalid_json_records_failed_usage(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    module = load_module(GUIDE_DRAFT_SCRIPT, "generate_guide_draft_failed_usage")
    module.USAGE_DB_PATH = tmp_path / "usage" / "generation_usage.sqlite3"
    trace_path = tmp_path / "procedure_trace.json"
    output_path = tmp_path / "generated" / "guide_draft.anthropic.json"
    trace_path.write_text(
        json.dumps({"schemaVersion": 1, "sessionId": "failed-session", "segments": []}),
        encoding="utf-8",
    )

    class FakeResponse:
        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, traceback) -> None:
            return None

        def read(self) -> bytes:
            return json.dumps(
                {
                    "anthropicResult": {
                        "content": [{"type": "text", "text": '{"title":"Unfinished"'}],
                        "usage": {"input_tokens": 2000, "output_tokens": 500},
                    },
                    "generationReport": {
                        "schemaVersion": 1,
                        "status": "succeeded",
                        "generatedAt": "2026-06-04T16:34:09Z",
                        "sessionId": "failed-session",
                        "title": "Failed guide generation",
                        "model": "claude-sonnet-4-6",
                        "provider": "anthropic",
                        "promptVersion": "guide-draft-v1",
                        "generationRunId": "failed-run-123",
                        "usage": {
                            "inputTokens": 2000,
                            "outputTokens": 500,
                            "totalTokens": 2500,
                            "estimatedCostUSD": 0.0135,
                        },
                    },
                }
            ).encode("utf-8")

    requests: list[str] = []

    def fake_urlopen(req, timeout):
        requests.append(str(req.full_url))
        return FakeResponse()

    monkeypatch.setenv("KCXDOC_REMOTE_API_BASE_URL", "https://kcxdocumentor-ai-dev.azurewebsites.net")
    monkeypatch.setenv("KCXDOC_REMOTE_API_BEARER_TOKEN", "token-123")
    monkeypatch.setattr(module.request, "urlopen", fake_urlopen)
    monkeypatch.setattr(sys, "argv", ["generate_guide_draft.py", str(trace_path), "--output", str(output_path)])

    result = module.main()

    assert result == 1
    assert not output_path.exists()
    failure_path = output_path.parent / "generation_failure.json"
    failure = json.loads(failure_path.read_text(encoding="utf-8"))
    assert failure["status"] == "failed"
    assert failure["sessionId"] == "failed-session"
    assert failure["generationRunId"] == "failed-run-123"
    assert failure["usage"] == {
        "inputTokens": 2000,
        "outputTokens": 500,
        "totalTokens": 2500,
        "estimatedCostUSD": 0.0135,
    }
    assert "invalid guide JSON" in failure["errorMessage"]
    assert requests == [
        "https://kcxdocumentor-ai-dev.azurewebsites.net/api/generate-draft",
        "https://kcxdocumentor-ai-dev.azurewebsites.net/api/usage-records",
    ]


def test_anthropic_http_error_writes_failure_without_traceback(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    module = load_module(GUIDE_DRAFT_SCRIPT, "generate_guide_draft_http_failure")
    module.USAGE_DB_PATH = tmp_path / "usage" / "generation_usage.sqlite3"
    trace_path = tmp_path / "procedure_trace.json"
    output_path = tmp_path / "generated" / "guide_draft.anthropic.json"
    trace_path.write_text(
        json.dumps({"schemaVersion": 1, "sessionId": "http-failed-session", "segments": []}),
        encoding="utf-8",
    )

    def raise_http_error(req, timeout):
        body = json.dumps(
            {
                "error": "Anthropic API request failed: HTTP 429 (rate_limit_error): Requests are temporarily limited.",
                "generationReport": {
                    "schemaVersion": 1,
                    "status": "failed",
                    "generatedAt": "2026-06-04T16:34:09Z",
                    "sessionId": "http-failed-session",
                    "title": "Failed guide generation",
                    "model": "claude-sonnet-4-6",
                    "provider": "anthropic",
                    "promptVersion": "guide-draft-v1",
                    "generationRunId": "http-failed-run",
                    "usage": {
                        "inputTokens": 0,
                        "outputTokens": 0,
                        "totalTokens": 0,
                        "estimatedCostUSD": 0,
                    },
                    "errorMessage": "Anthropic API request failed: HTTP 429 (rate_limit_error): Requests are temporarily limited.",
                },
            }
        ).encode("utf-8")
        raise module.error.HTTPError(str(req.full_url), 429, "Too Many Requests", {}, io.BytesIO(body))

    monkeypatch.setenv("KCXDOC_REMOTE_API_BASE_URL", "https://kcxdocumentor-ai-dev.azurewebsites.net")
    monkeypatch.setenv("KCXDOC_REMOTE_API_BEARER_TOKEN", "token-123")
    monkeypatch.setattr(module.request, "urlopen", raise_http_error)
    monkeypatch.setattr(sys, "argv", ["generate_guide_draft.py", str(trace_path), "--output", str(output_path)])

    result = module.main()

    assert result == 1
    failure_path = output_path.parent / "generation_failure.json"
    failure = json.loads(failure_path.read_text(encoding="utf-8"))
    assert failure["status"] == "failed"
    assert failure["sessionId"] == "http-failed-session"
    assert failure["usage"]["totalTokens"] == 0
    assert "HTTP 429" in failure["errorMessage"]
    assert "rate_limit_error" in failure["errorMessage"]
    assert "Traceback" not in failure["errorMessage"]


def test_anthropic_network_error_writes_failure_without_traceback(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    module = load_module(GUIDE_DRAFT_SCRIPT, "generate_guide_draft_network_failure")
    module.USAGE_DB_PATH = tmp_path / "usage" / "generation_usage.sqlite3"
    trace_path = tmp_path / "procedure_trace.json"
    output_path = tmp_path / "generated" / "guide_draft.anthropic.json"
    trace_path.write_text(
        json.dumps({"schemaVersion": 1, "sessionId": "network-failed-session", "segments": []}),
        encoding="utf-8",
    )

    def raise_network_error(req, timeout):
        raise module.error.URLError("connection reset")

    monkeypatch.setenv("KCXDOC_REMOTE_API_BASE_URL", "https://kcxdocumentor-ai-dev.azurewebsites.net")
    monkeypatch.setenv("KCXDOC_REMOTE_API_BEARER_TOKEN", "token-123")
    monkeypatch.setattr(module.request, "urlopen", raise_network_error)
    monkeypatch.setattr(sys, "argv", ["generate_guide_draft.py", str(trace_path), "--output", str(output_path)])

    result = module.main()

    assert result == 1
    failure = json.loads((output_path.parent / "generation_failure.json").read_text(encoding="utf-8"))
    assert failure["status"] == "failed"
    assert failure["sessionId"] == "network-failed-session"
    assert failure["usage"]["totalTokens"] == 0
    assert "failed before a complete response" in failure["errorMessage"]
    assert "Traceback" not in failure["errorMessage"]


def test_app_server_summarizes_generation_usage(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    module = load_module(APP_SERVER_SCRIPT, "app_server_generation_summary")
    processed_root = tmp_path / "processed"
    generated_root = tmp_path / "generated"
    session_dir = processed_root / "demo-session"
    generated_dir = generated_root / "demo-session"
    session_dir.mkdir(parents=True)
    generated_dir.mkdir(parents=True)
    (session_dir / "manifest.json").write_text(json.dumps({"createdUtc": "2026-06-04T16:00:00Z"}), encoding="utf-8")
    (session_dir / "procedure_trace.json").write_text(
        json.dumps(
            {
                "sessionId": "demo-session",
                "recording": {"targetApplication": "Blink Rx", "durationSeconds": 1241.601},
                "segments": [],
            }
        ),
        encoding="utf-8",
    )
    (generated_dir / "generation_report.json").write_text(
        json.dumps(
            {
                "schemaVersion": 1,
                "title": "Blink Rx Integration: Workflow Guide",
                "generatedAt": "2026-06-04T16:34:09Z",
                "model": "claude-sonnet-4-6",
                "provider": "anthropic",
                "promptVersion": "guide-draft-v1",
                "usage": {
                    "inputTokens": 12847,
                    "outputTokens": 3201,
                    "totalTokens": 16048,
                    "estimatedCostUSD": 0.086556,
                },
            }
        ),
        encoding="utf-8",
    )
    monkeypatch.setattr(module, "PROCESSED_ROOT", processed_root)
    monkeypatch.setattr(module, "GENERATED_ROOT", generated_root)
    monkeypatch.setattr(module, "WORKSPACE", tmp_path)

    session = module.read_session(session_dir)

    assert session["generation"] == {
        "title": "Blink Rx Integration: Workflow Guide",
        "model": "claude-sonnet-4-6",
        "provider": "anthropic",
        "promptVersion": "guide-draft-v1",
        "generatedAt": "2026-06-04T16:34:09Z",
        "inputTokens": 12847,
        "outputTokens": 3201,
        "totalTokens": 16048,
        "estimatedCostUSD": 0.086556,
        "pageCount": 0,
        "status": "succeeded",
        "errorMessage": "",
    }


def test_process_recording_accepts_sidecar_transcript_without_media_tools(tmp_path: Path) -> None:
    recording = tmp_path / "sample.mp4"
    transcript = tmp_path / "sample-transcript.txt"
    output_root = tmp_path / "processed"
    session_id = "sidecar-demo"
    recording.write_bytes(b"not a real video")
    transcript.write_text(
        "Open the customer record. Click Save after reviewing the address. Confirm the success message.",
        encoding="utf-8",
    )

    result = subprocess.run(
        [
            sys.executable,
            str(PROCESS_RECORDING_SCRIPT),
            str(recording),
            "--no-media-tools",
            "--transcript",
            str(transcript),
            "--output-root",
            str(output_root),
            "--session-id",
            session_id,
            "--target-application",
            "Enterprise Rx",
            "--assume-duration-seconds",
            "120",
            "--segment-seconds",
            "60",
            "--max-frames",
            "4",
        ],
        cwd=ROOT,
        check=False,
        text=True,
        capture_output=True,
    )

    assert result.returncode == 0, result.stdout + result.stderr
    session_dir = output_root / session_id
    transcript_payload = json.loads((session_dir / "transcript.json").read_text(encoding="utf-8"))
    trace_payload = json.loads((session_dir / "procedure_trace.json").read_text(encoding="utf-8"))

    assert transcript_payload["source"] == "sidecar-transcript"
    assert len(transcript_payload["segments"]) == 2
    assert all(segment["source"] == "sidecar-transcript" for segment in transcript_payload["segments"])
    assert "Click Save" in " ".join(segment["text"] for segment in transcript_payload["segments"])
    assert trace_payload["recording"]["targetApplication"] == "Enterprise Rx"
    assert len(trace_payload["segments"]) == 2
    assert trace_payload["segments"][0]["confidence"]["transcript"] >= 0.7
    assert trace_payload["segments"][0]["confidence"]["ocr"] < 0.7
    assert trace_payload["segments"][0]["confidence"]["frameSelection"] < 0.7
    assert trace_payload["segments"][0]["confidence"]["needsHumanReview"] is True


def test_tesseract_tsv_parser_groups_words_into_text_blocks() -> None:
    module = load_module(PROCESS_RECORDING_SCRIPT, "process_recording_ocr_parser")
    raw_tsv = "\n".join(
        [
            "level\tpage_num\tblock_num\tpar_num\tline_num\tword_num\tleft\ttop\twidth\theight\tconf\ttext",
            "5\t1\t1\t1\t1\t1\t10\t20\t42\t12\t92\tPatient",
            "5\t1\t1\t1\t1\t2\t58\t20\t38\t12\t88\tProfile",
            "5\t1\t1\t1\t2\t1\t10\t42\t31\t12\t81\tSave",
            "5\t1\t1\t1\t2\t2\t48\t42\t44\t12\t79\tChanges",
            "5\t1\t2\t1\t1\t1\t10\t80\t20\t12\t-1\t",
        ]
    )

    blocks = module.parse_tesseract_tsv(raw_tsv)

    assert blocks == [
        {
            "text": "Patient Profile",
            "confidence": 0.9,
            "bounds": {"left": 10, "top": 20, "width": 86, "height": 12},
        },
        {
            "text": "Save Changes",
            "confidence": 0.8,
            "bounds": {"left": 10, "top": 42, "width": 82, "height": 12},
        },
    ]


def test_build_ocr_runs_tesseract_and_stores_confidence(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    module = load_module(PROCESS_RECORDING_SCRIPT, "process_recording_ocr_builder")
    session_dir = tmp_path / "session"
    frame_path = session_dir / "frames" / "candidates" / "frame-0001.png"
    frame_path.parent.mkdir(parents=True)
    frame_path.write_bytes(b"fake image bytes")
    frame_scores = {
        "frames": [
            {
                "id": "frame-0001",
                "timestampSeconds": 12.0,
                "timestamp": "00:00:12.000",
                "path": "frames/candidates/frame-0001.png",
                "created": True,
            }
        ]
    }
    args = type(
        "Args",
        (),
        {
            "target_application": "Blink Rx",
            "ocr_language": "eng",
            "ocr_psm": "11",
            "ocr_timeout_seconds": 20.0,
        },
    )()

    def fake_run_command(cmd: list[str], timeout_seconds: float) -> dict[str, object]:
        assert cmd[:3] == ["/usr/bin/tesseract", str(frame_path), "stdout"]
        assert ["-l", "eng"] == cmd[3:5]
        assert ["--psm", "11"] == cmd[5:7]
        assert timeout_seconds == 20.0
        return {
            "returnCode": 0,
            "stdout": "\n".join(
                [
                    "level\tpage_num\tblock_num\tpar_num\tline_num\tword_num\tleft\ttop\twidth\theight\tconf\ttext",
                    "5\t1\t1\t1\t1\t1\t10\t20\t42\t12\t90\tRefill",
                    "5\t1\t1\t1\t1\t2\t58\t20\t38\t12\t86\tRequest",
                ]
            ),
            "stderr": "",
        }

    monkeypatch.setattr(module, "run_command", fake_run_command)

    payload = module.build_ocr(
        frame_scores,
        session_dir,
        module.Tooling(ffprobe=None, ffmpeg=None, whisper=None, tesseract="/usr/bin/tesseract"),
        args,
    )

    assert payload["source"] == "tesseract"
    assert payload["frames"][0]["source"] == "tesseract"
    assert payload["frames"][0]["combinedText"] == "Refill Request"
    assert payload["frames"][0]["confidence"] == 0.88
    assert payload["frames"][0]["textBlocks"][0]["bounds"] == {"left": 10, "top": 20, "width": 86, "height": 12}


def test_candidate_image_penalizes_teams_title_card_ocr() -> None:
    module = load_module(PROCESS_RECORDING_SCRIPT, "process_recording_ocr_frame_scoring")
    frame = {
        "id": "frame-title",
        "timestamp": "00:01:00.000",
        "timestampSeconds": 60.0,
        "path": "frames/candidates/frame-title.png",
        "webPath": "frames/candidates/frame-title.png",
        "score": 0.82,
        "created": True,
        "selectionReason": "Extracted at regular interval for prototype review.",
    }
    ocr_frame = {
        "source": "tesseract",
        "confidence": 0.94,
        "combinedText": "Microsoft Teams Newleaf and General Pharmacy Industry Training Sessions 2025-12-05 15:01 UTC Recorded by Tina Drake",
    }
    segment = {"text": "Open the refill request screen and review the patient profile."}

    image = module.build_candidate_image(frame, ocr_frame, segment, "Blink Rx")

    assert image["ocrNonApplication"] is True
    assert image["frameEvidenceScore"] <= 0.34
    assert "Penalized" in image["reason"]


def test_candidate_image_boosts_application_ocr_overlap() -> None:
    module = load_module(PROCESS_RECORDING_SCRIPT, "process_recording_ocr_app_scoring")
    frame = {
        "id": "frame-app",
        "timestamp": "00:02:00.000",
        "timestampSeconds": 120.0,
        "path": "frames/candidates/frame-app.png",
        "webPath": "frames/candidates/frame-app.png",
        "score": 0.72,
        "created": True,
        "selectionReason": "Extracted at regular interval for prototype review.",
    }
    ocr_frame = {
        "source": "tesseract",
        "confidence": 0.64,
        "combinedText": "Blink Mock UI Refill Request Patient Profile Submit Create Request",
    }
    segment = {"text": "Create the refill request from the patient profile and submit it."}

    image = module.build_candidate_image(frame, ocr_frame, segment, "Blink Rx")

    assert image["ocrNonApplication"] is False
    assert image["ocrRelevanceScore"] >= 0.4
    assert image["frameEvidenceScore"] > 0.5


def test_candidate_image_penalizes_supporting_tool_when_segment_is_application_workflow() -> None:
    module = load_module(PROCESS_RECORDING_SCRIPT, "process_recording_supporting_tool_scoring")
    frame = {
        "id": "frame-json",
        "timestamp": "00:02:00.000",
        "timestampSeconds": 120.0,
        "path": "frames/candidates/frame-json.png",
        "webPath": "frames/candidates/frame-json.png",
        "score": 0.83,
        "created": True,
        "selectionReason": "Extracted at regular interval for prototype review.",
    }
    ocr_frame = {
        "source": "tesseract",
        "confidence": 0.82,
        "combinedText": "C:\\BlinkMockData\\Requests.json Notepad++ Administrator refill initiation request",
    }
    segment = {"text": "Complete the refill request in Blink Rx and verify the prescription screen."}

    image = module.build_candidate_image(frame, ocr_frame, segment, "Blink Rx")

    assert image["ocrSupportingTool"] is True
    assert image["frameEvidenceScore"] <= 0.46
    assert "supporting tool" in image["reason"]


def test_candidate_image_allows_supporting_tool_for_mock_json_segment() -> None:
    module = load_module(PROCESS_RECORDING_SCRIPT, "process_recording_supporting_tool_allowed")
    frame = {
        "id": "frame-json",
        "timestamp": "00:02:00.000",
        "timestampSeconds": 120.0,
        "path": "frames/candidates/frame-json.png",
        "webPath": "frames/candidates/frame-json.png",
        "score": 0.83,
        "created": True,
        "selectionReason": "Extracted at regular interval for prototype review.",
    }
    ocr_frame = {
        "source": "tesseract",
        "confidence": 0.82,
        "combinedText": "C:\\BlinkMockData\\Requests.json Notepad++ Administrator refill initiation request",
    }
    segment = {"text": "Update the mock request JSON value before rerunning the refill test."}

    image = module.build_candidate_image(frame, ocr_frame, segment, "Blink Rx")

    assert image["ocrSupportingTool"] is True
    assert image["frameEvidenceScore"] > 0.46


def test_generate_draft_prefers_application_frame_over_clean_teams_title_card() -> None:
    module = load_module(GUIDE_DRAFT_SCRIPT, "generate_guide_draft_ocr_sorting")
    teams_title = {
        "frameId": "teams-title",
        "path": "frames/candidates/teams-title.png",
        "timestampSeconds": 60.0,
        "score": 0.9,
        "confidence": 0.34,
        "frameEvidenceScore": 0.34,
        "ocrConfidence": 0.94,
        "ocrNonApplication": True,
        "created": True,
        "reviewStatus": "pending",
    }
    app_frame = {
        "frameId": "app-frame",
        "path": "frames/candidates/app-frame.png",
        "timestampSeconds": 90.0,
        "score": 0.72,
        "confidence": 0.58,
        "frameEvidenceScore": 0.58,
        "ocrConfidence": 0.64,
        "ocrNonApplication": False,
        "created": True,
        "reviewStatus": "pending",
    }
    draft = {"steps": [{"instruction": "Create the refill request.", "sourceSegments": ["seg-0001"]}]}
    trace = {
        "sessionId": "demo-session",
        "segments": [{"id": "seg-0001", "candidateImages": [teams_title, app_frame]}],
    }

    enriched = module.attach_screenshot_references(draft, trace)

    assert enriched["steps"][0]["screenshotRef"] == "app-frame"


def test_generate_draft_prefers_application_frame_over_supporting_tool() -> None:
    module = load_module(GUIDE_DRAFT_SCRIPT, "generate_guide_draft_supporting_tool_sorting")
    supporting_tool = {
        "frameId": "json-frame",
        "path": "frames/candidates/json-frame.png",
        "timestampSeconds": 60.0,
        "score": 0.9,
        "confidence": 0.46,
        "frameEvidenceScore": 0.46,
        "ocrSupportingTool": True,
        "created": True,
        "reviewStatus": "pending",
    }
    app_frame = {
        "frameId": "app-frame",
        "path": "frames/candidates/app-frame.png",
        "timestampSeconds": 80.0,
        "score": 0.7,
        "confidence": 0.5,
        "frameEvidenceScore": 0.5,
        "created": True,
        "reviewStatus": "pending",
    }
    draft = {"steps": [{"instruction": "Complete the refill request.", "sourceSegments": ["seg-0001"]}]}
    trace = {
        "sessionId": "demo-session",
        "segments": [{"id": "seg-0001", "candidateImages": [supporting_tool, app_frame]}],
    }

    enriched = module.attach_screenshot_references(draft, trace)

    assert enriched["steps"][0]["screenshotRef"] == "app-frame"


def test_generate_draft_prefers_unique_sharp_frame_over_blurry_duplicate() -> None:
    module = load_module(GUIDE_DRAFT_SCRIPT, "generate_guide_draft_visual_quality_sorting")
    duplicate = {
        "frameId": "duplicate",
        "path": "frames/candidates/duplicate.png",
        "timestampSeconds": 60.0,
        "score": 0.9,
        "confidence": 0.72,
        "frameEvidenceScore": 0.72,
        "visualQualityScore": 0.3,
        "dedupeState": "near-duplicate",
        "blurState": "blurry",
        "ocrClass": "application",
        "created": True,
        "reviewStatus": "pending",
    }
    sharp = {
        "frameId": "sharp",
        "path": "frames/candidates/sharp.png",
        "timestampSeconds": 80.0,
        "score": 0.68,
        "confidence": 0.7,
        "frameEvidenceScore": 0.7,
        "visualQualityScore": 0.9,
        "dedupeState": "unique",
        "blurState": "sharp",
        "ocrClass": "application",
        "created": True,
        "reviewStatus": "pending",
    }
    draft = {"steps": [{"instruction": "Complete the refill request.", "sourceSegments": ["seg-0001"]}]}
    trace = {
        "sessionId": "demo-session",
        "segments": [{"id": "seg-0001", "candidateImages": [duplicate, sharp]}],
    }

    enriched = module.attach_screenshot_references(draft, trace)

    assert enriched["steps"][0]["screenshotRef"] == "sharp"


def test_caption_parser_ignores_teams_cue_ids_and_voice_tags() -> None:
    module = load_module(PROCESS_RECORDING_SCRIPT, "process_recording_caption_parser")
    raw_text = """WEBVTT

724b054c-5275-4678-b67e-a50dcb6dc1a2/45-0
00:00:12.796 --> 00:00:18.129
<v Tina Drake>And today we're going to cover the last
of Blink. Um,</v>

724b054c-5275-4678-b67e-a50dcb6dc1a2/45-1
00:00:18.129 --> 00:00:23.956
<v Tina Drake>so you know she's going to run through
what was remaining.</v>
"""

    segments = module.parse_caption_transcript(raw_text)

    assert len(segments) == 2
    assert segments[0]["text"] == "And today we're going to cover the last of Blink. Um,"
    assert segments[1]["text"] == "so you know she's going to run through what was remaining."
    assert "724b054c" not in " ".join(segment["text"] for segment in segments)
    assert "<v" not in " ".join(segment["text"] for segment in segments)


def test_timed_transcript_segments_are_bucketed_across_recording() -> None:
    module = load_module(PROCESS_RECORDING_SCRIPT, "process_recording_timed_buckets")

    chunks = module.split_source_segments(
        [
            {"text": "Open the first screen.", "startSeconds": 5.0, "endSeconds": 9.0, "confidence": 0.9},
            {"text": "Review the middle workflow.", "startSeconds": 75.0, "endSeconds": 90.0, "confidence": 0.8},
            {"text": "Confirm the final result.", "startSeconds": 125.0, "endSeconds": 135.0, "confidence": 0.7},
        ],
        count=3,
        duration=180.0,
        segment_seconds=60.0,
    )

    assert [chunk["text"] for chunk in chunks] == [
        "Open the first screen.",
        "Review the middle workflow.",
        "Confirm the final result.",
    ]
    assert chunks[0]["startSeconds"] == 0.0
    assert chunks[2]["endSeconds"] == 180.0


def test_compare_transcripts_reports_overlap_and_examples() -> None:
    module = load_module(COMPARE_TRANSCRIPTS_SCRIPT, "compare_transcripts_metrics")

    report = module.compare_transcripts(
        {
            "source": "sidecar-transcript",
            "segments": [
                {"id": "tx-0001", "startSeconds": 0, "endSeconds": 60, "confidence": 0.78, "text": "Open the refill request."},
                {"id": "tx-0002", "startSeconds": 60, "endSeconds": 120, "confidence": 0.78, "text": "Click submit to save."},
            ],
        },
        {
            "source": "local-whisper",
            "segments": [
                {"id": "tx-0001", "startSeconds": 0, "endSeconds": 60, "confidence": 0.88, "text": "Open the refill request."},
                {"id": "tx-0002", "startSeconds": 60, "endSeconds": 120, "confidence": 0.6, "text": "Click save."},
            ],
        },
        example_count=1,
    )

    assert report["reference"]["source"] == "sidecar-transcript"
    assert report["candidate"]["source"] == "local-whisper"
    assert report["comparison"]["wordOverlap"] > 0.5
    assert report["comparison"]["averageAlignedSimilarity"] > 0.7
    assert len(report["examples"]) == 1


def test_whisper_json_segments_include_offsets_and_confidence() -> None:
    module = load_module(PROCESS_RECORDING_SCRIPT, "process_recording_whisper_parser")
    payload = {
        "result": {"language": "en"},
        "transcription": [
            {
                "offsets": {"from": 1200, "to": 5400},
                "text": " Click Save to finish the workflow.",
                "tokens": [
                    {"text": "[_BEG_]", "p": 0.1},
                    {"text": " Click", "p": 0.9},
                    {"text": " Save", "p": 0.8},
                    {"text": ".", "p": 0.7},
                ],
            }
        ],
    }

    segments = module.parse_whisper_transcript(payload)

    assert segments == [
        {
            "id": "whisper-0001",
            "text": "Click Save to finish the workflow.",
            "startSeconds": 1.2,
            "endSeconds": 5.4,
            "confidence": 0.8,
            "speaker": "Speaker 1",
        }
    ]


def test_build_transcript_prefers_local_whisper_when_no_sidecar() -> None:
    module = load_module(PROCESS_RECORDING_SCRIPT, "process_recording_whisper_transcript")
    transcript = module.build_transcript(
        metadata={"durationSeconds": 30.0},
        sidecar_transcript=None,
        local_transcript={
            "source": "local-whisper",
            "name": "whisper-transcript.json",
            "path": "audio/whisper-transcript.json",
            "format": "json",
            "model": "models/whisper/ggml-base.en.bin",
            "language": "en",
            "error": None,
            "segments": [
                {
                    "text": "Open the order screen and select New.",
                    "startSeconds": 0.0,
                    "endSeconds": 12.0,
                    "confidence": 0.87,
                    "speaker": "Speaker 1",
                }
            ],
        },
        segment_seconds=60.0,
        target_application="Enterprise Rx",
    )

    assert transcript["source"] == "local-whisper"
    assert transcript["sourceTranscript"]["source"] == "local-whisper"
    assert transcript["sourceTranscript"]["path"] == "audio/whisper-transcript.json"
    assert transcript["segments"][0]["source"] == "local-whisper"
    assert transcript["segments"][0]["confidence"] == 0.87
    assert "Open the order screen" in transcript["segments"][0]["text"]


def test_teams_recording_profile_skips_intro_and_crops_frames() -> None:
    module = load_module(PROCESS_RECORDING_SCRIPT, "process_recording_teams_profile")
    args = type("Args", (), {"source_profile": "teams-recording", "skip_start_seconds": None, "frame_crop_filter": None})()

    assert module.effective_skip_start_seconds(args) == 60.0
    assert module.build_frame_crop_filter(args).startswith("crop=")
    assert module.planned_frame_timestamps(300.0, 30.0, 3, start_seconds=60.0) == [60.0, 90.0, 120.0]


def test_score_frames_marks_near_duplicate_frames(monkeypatch: pytest.MonkeyPatch) -> None:
    module = load_module(PROCESS_RECORDING_SCRIPT, "process_recording_visual_dedupe")
    frames = [
        {"id": "frame-1", "timestampSeconds": 10.0, "timestamp": "00:00:10.000", "created": True, "path": "frame-1.png"},
        {"id": "frame-2", "timestampSeconds": 20.0, "timestamp": "00:00:20.000", "created": True, "path": "frame-2.png"},
    ]

    def fake_quality(frame, session_dir):
        return {
            "qualityScore": 0.8,
            "sharpnessScore": 0.8,
            "exposureScore": 0.8,
            "blurState": "sharp",
            "exposureState": "usable",
            "averageHash": "ffff0000ffff0000",
            "visualScoringAvailable": True,
        }

    monkeypatch.setattr(module, "evaluate_frame_visual_quality", fake_quality)

    scored = module.score_frames(frames, {"durationSeconds": 120.0}, 10.0)

    assert scored["frames"][0]["qualitySignals"]["dedupeState"] == "unique"
    assert scored["frames"][1]["qualitySignals"]["dedupeState"] == "near-duplicate"
    assert scored["frames"][1]["qualitySignals"]["duplicateOfFrameId"] == "frame-1"
    assert scored["frames"][1]["score"] < scored["frames"][0]["score"]


def test_nearest_frames_prefers_unique_sharp_candidate_over_closer_duplicate() -> None:
    module = load_module(PROCESS_RECORDING_SCRIPT, "process_recording_nearest_quality")
    frames = [
        {
            "id": "duplicate",
            "timestampSeconds": 60.0,
            "score": 0.9,
            "qualitySignals": {"dedupeState": "near-duplicate", "blurState": "blurry"},
        },
        {
            "id": "sharp",
            "timestampSeconds": 70.0,
            "score": 0.7,
            "qualitySignals": {"dedupeState": "unique", "blurState": "sharp"},
        },
    ]

    selected = module.nearest_frames(frames, 55.0, 75.0, 1)

    assert selected[0]["id"] == "sharp"


def test_ocr_surface_classifier_identifies_application_and_person_frames() -> None:
    module = load_module(PROCESS_RECORDING_SCRIPT, "process_recording_ocr_classifier")

    app = module.classify_ocr_surface("Blink Mock UI MANOJI'S PHARMACY Data Entry Search for Rx and Refills")
    person = module.classify_ocr_surface("Vibindas Asokakumar (UST, IN)")

    assert app["ocrClass"] == "application"
    assert app["appOcrScore"] > 0
    assert person["ocrClass"] == "non-application"


def test_frame_recommendation_groups_prune_weak_non_application_candidates() -> None:
    module = load_module(PROCESS_RECORDING_SCRIPT, "process_recording_frame_groups")
    images = [
        {
            "frameId": "app",
            "ocrClass": "application",
            "frameEvidenceScore": 0.74,
            "visualQualityScore": 0.8,
            "blurState": "sharp",
            "timestampSeconds": 20.0,
        },
        {
            "frameId": "backup",
            "ocrClass": "application",
            "frameEvidenceScore": 0.5,
            "visualQualityScore": 0.7,
            "blurState": "sharp",
            "timestampSeconds": 30.0,
        },
        {
            "frameId": "teams",
            "ocrClass": "non-application",
            "ocrNonApplication": True,
            "frameEvidenceScore": 0.7,
            "visualQualityScore": 0.8,
            "timestampSeconds": 10.0,
        },
    ]

    grouped = module.assign_frame_recommendation_groups(images)
    by_id = {image["frameId"]: image for image in grouped}

    assert by_id["app"]["recommendationGroup"] == "recommended"
    assert by_id["backup"]["recommendationGroup"] == "alternate"
    assert by_id["teams"]["recommendationGroup"] == "system-rejected"
    assert by_id["teams"]["penalties"]


def test_segment_quality_creates_screenshot_gap_when_no_recommended_frame() -> None:
    module = load_module(PROCESS_RECORDING_SCRIPT, "process_recording_segment_quality")
    segment = {
        "id": "seg-0001",
        "start": "00:00:10.000",
        "end": "00:00:20.000",
        "startSeconds": 10.0,
        "endSeconds": 20.0,
    }
    quality = module.build_segment_quality(
        segment,
        {"transcript": 0.92, "ocr": 0.4, "frameSelection": 0.3, "overall": 0.6, "needsHumanReview": True},
        [{"frameId": "teams", "recommendationGroup": "system-rejected", "ocrNonApplication": True}],
        [],
    )

    assert quality["qualityLabel"] == "needs-review"
    assert quality["reviewPriority"] == "high"
    assert quality["screenshotGap"]["needsBetterScreenshot"] is True
    assert any(label["id"] == "missing-recommended-screenshot" for label in quality["qualityLabels"])


def test_detect_recording_content_type_identifies_application_workflow() -> None:
    module = load_module(PROCESS_RECORDING_SCRIPT, "process_recording_content_type")
    segments = [
        {
            "speakerText": "Click save, open the refill request, and submit it.",
            "actionHints": ["click", "submit"],
            "candidateImages": [{"ocrClass": "application"}],
        },
        {
            "speakerText": "Select the pharmacy profile and save the plan.",
            "actionHints": ["select", "save"],
            "candidateImages": [{"ocrClass": "application"}],
        },
    ]

    classification = module.detect_recording_content_type(segments, {"source": "sidecar-transcript"})

    assert classification["type"] == "application-workflow"
    assert classification["confidence"] > 0.7


def test_generate_draft_attaches_screenshot_references_to_model_steps() -> None:
    module = load_module(GUIDE_DRAFT_SCRIPT, "generate_guide_draft_screenshots")
    draft = {
        "sections": [
            {
                "title": "Approve Fax",
                "steps": [
                    {
                        "instruction": "Select the fax row.",
                        "sourceSegments": ["seg-0002"],
                    }
                ],
            }
        ]
    }
    trace = {
        "sessionId": "demo-session",
        "segments": [
            {
                "id": "seg-0002",
                "candidateImages": [
                    {
                        "frameId": "frame-0002",
                        "path": "frames/candidates/frame-0002.jpg",
                        "timestamp": "00:01:00.000",
                        "timestampSeconds": 60.0,
                        "score": 0.81,
                        "frameEvidenceScore": 0.72,
                        "visualQualityScore": 0.88,
                        "ocrClass": "application",
                        "created": True,
                        "reviewStatus": "pending",
                    }
                ],
            }
        ],
    }

    enriched = module.attach_screenshot_references(draft, trace)
    step = enriched["sections"][0]["steps"][0]

    assert step["screenshotRef"] == "frame-0002"
    assert step["selectedScreenshot"]["frameId"] == "frame-0002"
    assert step["selectedScreenshot"]["frameEvidenceScore"] == 0.72
    assert step["selectedScreenshot"]["visualQualityScore"] == 0.88
    assert step["selectedScreenshot"]["ocrClass"] == "application"
    assert step["selectedScreenshot"]["path"].endswith("samples/processed/demo-session/frames/candidates/frame-0002.jpg")


def test_generate_draft_distributes_screenshots_when_model_omits_source_segments() -> None:
    module = load_module(GUIDE_DRAFT_SCRIPT, "generate_guide_draft_screenshot_distribution")
    draft = {
        "steps": [
            {"instruction": "First action."},
            {"instruction": "Second action."},
            {"instruction": "Third action."},
        ]
    }
    trace = {
        "sessionId": "demo-session",
        "segments": [
            {
                "id": "seg-0001",
                "candidateImages": [
                    {
                        "frameId": "frame-0001",
                        "path": "frames/candidates/frame-0001.jpg",
                        "timestampSeconds": 60.0,
                        "score": 0.7,
                        "created": True,
                    }
                ],
            },
            {
                "id": "seg-0002",
                "candidateImages": [
                    {
                        "frameId": "frame-0002",
                        "path": "frames/candidates/frame-0002.jpg",
                        "timestampSeconds": 120.0,
                        "score": 0.7,
                        "created": True,
                    }
                ],
            },
            {
                "id": "seg-0003",
                "candidateImages": [
                    {
                        "frameId": "frame-0003",
                        "path": "frames/candidates/frame-0003.jpg",
                        "timestampSeconds": 180.0,
                        "score": 0.7,
                        "created": True,
                    }
                ],
            },
        ],
    }

    enriched = module.attach_screenshot_references(draft, trace)

    assert [step["screenshotRef"] for step in enriched["steps"]] == ["frame-0001", "frame-0002", "frame-0003"]


def test_generate_draft_avoids_reusing_screenshot_when_alternative_exists() -> None:
    module = load_module(GUIDE_DRAFT_SCRIPT, "generate_guide_draft_no_reuse")
    draft = {
        "steps": [
            {"instruction": "First action.", "sourceSegments": ["seg-0001"]},
            {"instruction": "Second action.", "sourceSegments": ["seg-0001"]},
        ]
    }
    trace = {
        "sessionId": "demo-session",
        "segments": [
            {
                "id": "seg-0001",
                "candidateImages": [
                    {
                        "frameId": "frame-best",
                        "path": "frames/candidates/frame-best.jpg",
                        "timestampSeconds": 60.0,
                        "frameEvidenceScore": 0.9,
                        "score": 0.9,
                        "created": True,
                    },
                    {
                        "frameId": "frame-second",
                        "path": "frames/candidates/frame-second.jpg",
                        "timestampSeconds": 80.0,
                        "frameEvidenceScore": 0.75,
                        "score": 0.75,
                        "created": True,
                    },
                ],
            }
        ],
    }

    enriched = module.attach_screenshot_references(draft, trace)

    assert [step["screenshotRef"] for step in enriched["steps"]] == ["frame-best", "frame-second"]


def test_generate_draft_applies_frame_review_file_to_trace(tmp_path: Path) -> None:
    module = load_module(GUIDE_DRAFT_SCRIPT, "generate_guide_draft_frame_review")
    trace_path = tmp_path / "procedure_trace.json"
    (tmp_path / "frame_scores.json").write_text(
        json.dumps(
            {
                "frames": [
                    {
                        "id": "review-frame-0001",
                        "path": "frames/candidates/review-frame-0001.png",
                        "webPath": "frames/candidates/review-frame-0001.png",
                        "timestamp": "00:02:00.000",
                        "timestampSeconds": 120.0,
                        "score": 0.9,
                        "created": True,
                        "selectionReason": "Added by reviewer.",
                        "frameEvidenceScore": 0.91,
                        "ocrConfidence": 0.88,
                        "ocrRelevanceScore": 0.8,
                        "ocrClass": "application",
                        "ocrText": "Submit Refill Request",
                        "contentType": "application",
                        "recommendationGroup": "recommended",
                        "selectionDecision": "recommended",
                        "recommendationReason": "Best manual screenshot candidate.",
                    }
                ]
            }
        ),
        encoding="utf-8",
    )
    (tmp_path / "frame_review.json").write_text(
        json.dumps(
            {
                "schemaVersion": 1,
                "frames": {
                    "frame-0001": {"status": "rejected", "note": "Teams title card."},
                    "review-frame-0001": {
                        "status": "approved",
                        "note": "Use this dialog.",
                        "assignedSegmentId": "seg-0001",
                        "addedByReviewer": True,
                    },
                },
            }
        ),
        encoding="utf-8",
    )
    trace = {
        "sessionId": "review-demo",
        "segments": [
            {
                "id": "seg-0001",
                "candidateImages": [
                    {
                        "frameId": "frame-0001",
                        "path": "frames/candidates/frame-0001.png",
                        "timestampSeconds": 30.0,
                        "score": 0.7,
                        "created": True,
                        "reviewStatus": "pending",
                    }
                ],
            }
        ],
    }
    trace_path.write_text(json.dumps(trace), encoding="utf-8")

    reviewed = module.apply_frame_review(trace, trace_path)
    images = reviewed["segments"][0]["candidateImages"]

    assert images[0]["reviewStatus"] == "rejected"
    assert images[0]["reviewNote"] == "Teams title card."
    assert images[1]["frameId"] == "review-frame-0001"
    assert images[1]["reviewStatus"] == "approved"
    assert images[1]["reviewNote"] == "Use this dialog."
    assert images[1]["ocrText"] == "Submit Refill Request"
    assert images[1]["ocrClass"] == "application"
    assert images[1]["recommendationGroup"] == "recommended"


@pytest.mark.parametrize(
    "relative_path",
    [
        "scripts/process_recording.py",
        "scripts/build_guide_docx.py",
        "scripts/generate_guide_draft.py",
    ],
)
def test_future_prototype_scripts_have_smoke_test_hooks(relative_path: str) -> None:
    script = ROOT / relative_path
    if not script.exists():
        pytest.skip(f"{relative_path} is not implemented yet")

    help_result = subprocess.run(
        [sys.executable, str(script), "--help"],
        cwd=ROOT,
        check=False,
        text=True,
        capture_output=True,
    )

    assert help_result.returncode == 0
    help_text = help_result.stdout + help_result.stderr
    assert ("--input" in help_text or "recording" in help_text or "trace" in help_text)
    assert ("--output" in help_text or "--output-root" in help_text)
    if relative_path.endswith("process_recording.py"):
        assert "--no-media-tools" in help_text
        assert "--whisper-model" in help_text
        assert "--no-local-stt" in help_text
