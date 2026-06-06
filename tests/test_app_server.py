from __future__ import annotations

import importlib.util
import inspect
import io
import json
import os
import sqlite3
import sys
from pathlib import Path
from typing import Any

import pytest
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
APP_SERVER = ROOT / "scripts" / "app_server.py"


@pytest.fixture(autouse=True)
def clear_remote_usage_api(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.delenv("KCXDOC_REMOTE_API_BASE_URL", raising=False)


def load_app_server():
    if not APP_SERVER.exists():
        pytest.skip("scripts/app_server.py is not implemented yet")

    spec = importlib.util.spec_from_file_location("kcx_app_server_test", APP_SERVER)
    assert spec and spec.loader
    module = importlib.util.module_from_spec(spec)
    sys.modules["kcx_app_server_test"] = module
    spec.loader.exec_module(module)
    return module


def find_callable(module: Any, *names: str):
    for name in names:
        candidate = getattr(module, name, None)
        if callable(candidate):
            return candidate
    pytest.skip(f"app server does not expose any of: {', '.join(names)}")


def find_optional_callable(module: Any, *names: str):
    for name in names:
        candidate = getattr(module, name, None)
        if callable(candidate):
            return candidate
    return None


class FakeHandler:
    def __init__(self) -> None:
        self.status: int | None = None
        self.headers: list[tuple[str, str]] = []
        self.wfile = io.BytesIO()

    def send_response(self, status: int) -> None:
        self.status = status

    def send_header(self, key: str, value: str) -> None:
        self.headers.append((key, value))

    def end_headers(self) -> None:
        pass


def invoke_with_supported_kwargs(func: Any, *args: Any, **kwargs: Any) -> Any:
    signature = inspect.signature(func)
    supported = {
        key: value
        for key, value in kwargs.items()
        if key in signature.parameters
    }
    return func(*args, **supported)


def test_json_response_helper_writes_minimal_http_response() -> None:
    module = load_app_server()
    helper = find_callable(module, "send_json", "write_json", "write_json_response")
    handler = FakeHandler()

    invoke_with_supported_kwargs(
        helper,
        handler,
        {"ok": True, "message": "ready"},
        status=202,
        status_code=202,
    )

    assert handler.status == 202
    headers = dict(handler.headers)
    assert headers.get("Content-Type", "").startswith("application/json")
    assert "no-store" in headers.get("Cache-Control", "").lower()
    payload = json.loads(handler.wfile.getvalue().decode("utf-8"))
    assert payload == {"ok": True, "message": "ready"}


def test_json_response_builder_is_utf8_bytes_when_exposed() -> None:
    module = load_app_server()
    builder = find_callable(module, "build_json_response", "json_response_bytes")

    response = builder({"status": "ok", "application": "Enterprise Rx"})

    assert isinstance(response, bytes)
    assert json.loads(response.decode("utf-8")) == {
        "status": "ok",
        "application": "Enterprise Rx",
    }


def test_auth_config_uses_local_environment(monkeypatch: pytest.MonkeyPatch) -> None:
    module = load_app_server()
    monkeypatch.setenv("KCXDOC_AUTH_ENABLED", "true")
    monkeypatch.setenv("KCXDOC_AUTH_TENANT_ID", "tenant-123")
    monkeypatch.setenv("KCXDOC_AUTH_CLIENT_ID", "client-456")
    monkeypatch.setenv("KCXDOC_AUTH_AUTHORITY", "https://login.microsoftonline.com/tenant-123")
    monkeypatch.setenv("KCXDOC_AUTH_REDIRECT_URI", "http://127.0.0.1:8765/")
    monkeypatch.setenv("KCXDOC_AUTH_POST_LOGOUT_REDIRECT_URI", "http://127.0.0.1:8765/")
    monkeypatch.setenv("KCXDOC_AUTH_SCOPES", "openid profile")

    config = module.get_auth_config()

    assert config == {
        "enabled": True,
        "tenantId": "tenant-123",
        "clientId": "client-456",
        "authority": "https://login.microsoftonline.com/tenant-123",
        "redirectUri": "http://127.0.0.1:8765/",
        "postLogoutRedirectUri": "http://127.0.0.1:8765/",
        "scopes": ["openid", "profile"],
    }


def test_auth_session_cookie_is_signed_and_expires(monkeypatch: pytest.MonkeyPatch) -> None:
    module = load_app_server()
    monkeypatch.setenv("KCXDOC_AUTH_ENABLED", "true")
    monkeypatch.setenv("KCXDOC_AUTH_TENANT_ID", "tenant-123")
    monkeypatch.setenv("KCXDOC_AUTH_CLIENT_ID", "client-456")
    monkeypatch.setattr(module, "AUTH_SESSION_SECRET", "unit-test-secret")
    monkeypatch.setattr(module, "time_now", lambda: 1000)

    cookie = module.build_auth_session_cookie(
        {
            "aud": "client-456",
            "iss": "https://login.microsoftonline.com/tenant-123/v2.0",
            "sub": "user-1",
            "exp": 1600,
        },
        600,
    )

    assert module.validate_auth_session_cookie(cookie) is True
    assert module.validate_auth_session_cookie(cookie.replace("a", "b", 1)) is False
    monkeypatch.setattr(module, "time_now", lambda: 1601)
    assert module.validate_auth_session_cookie(cookie) is False


def test_safe_path_helper_accepts_children_and_rejects_traversal(tmp_path: Path) -> None:
    module = load_app_server()
    safe_path = find_callable(module, "safe_join", "resolve_safe_path", "safe_resolve")
    root = tmp_path / "workspace"
    root.mkdir()
    (root / "samples").mkdir()
    (root / "samples" / "demo.mp4").write_bytes(b"placeholder")

    resolved = safe_path(root, "samples/demo.mp4")

    assert Path(resolved).resolve() == (root / "samples" / "demo.mp4").resolve()
    with pytest.raises((PermissionError, ValueError)):
        safe_path(root, "../outside.txt")
    with pytest.raises((PermissionError, ValueError)):
        safe_path(root, str(tmp_path / "outside.txt"))


def test_app_state_lists_processed_sessions(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    module = load_app_server()
    processed_root = tmp_path / "samples" / "processed"
    session_dir = processed_root / "session-a"
    session_dir.mkdir(parents=True)
    (session_dir / "manifest.json").write_text(
        json.dumps(
            {
                "schemaVersion": 1,
                "sessionId": "session-a",
                "sourceFile": "samples/raw/example.mp4",
                "outputs": {"procedureTrace": "procedure_trace.json"},
            },
        ),
        encoding="utf-8",
    )
    monkeypatch.setattr(module, "PROCESSED_ROOT", processed_root)
    monkeypatch.setattr(module, "GENERATED_ROOT", tmp_path / "artifacts" / "generated")

    sessions = module.list_sessions()

    assert len(sessions) == 1
    assert sessions[0]["sessionId"] == "session-a"
    assert sessions[0]["segmentCount"] is None


def test_session_modified_time_uses_generated_artifacts(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    module = load_app_server()
    processed_root = tmp_path / "samples" / "processed"
    generated_root = tmp_path / "artifacts" / "generated"
    session_dir = processed_root / "session-a"
    generated_dir = generated_root / "session-a"
    session_dir.mkdir(parents=True)
    generated_dir.mkdir(parents=True)
    (session_dir / "manifest.json").write_text(json.dumps({"schemaVersion": 1, "sessionId": "session-a"}), encoding="utf-8")
    docx = generated_dir / "user_guide.anthropic.docx"
    docx.write_bytes(b"docx")
    os.utime(session_dir, (1000, 1000))
    os.utime(docx, (2000, 2000))
    monkeypatch.setattr(module, "PROCESSED_ROOT", processed_root)
    monkeypatch.setattr(module, "GENERATED_ROOT", generated_root)

    session = module.summarize_session(session_dir)

    assert session["modifiedUtc"] == module.utc_from_timestamp(2000)


def test_usage_summary_aggregates_generation_reports_by_range(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    module = load_app_server()
    generated_root = tmp_path / "artifacts" / "generated"
    first = generated_root / "session-a"
    second = generated_root / "session-b"
    first.mkdir(parents=True)
    second.mkdir(parents=True)
    (first / "generation_report.json").write_text(
        json.dumps(
            {
                "generatedAt": "2026-06-04T16:34:09Z",
                "sessionId": "session-a",
                "title": "Guide A",
                "model": "claude-sonnet-4-6",
                "pageCount": 4,
                "generatedBy": {
                    "oid": "user-a",
                    "name": "Analyst One",
                    "username": "analyst.one@example.com",
                },
                "usage": {
                    "inputTokens": 1000,
                    "outputTokens": 200,
                    "totalTokens": 1200,
                    "estimatedCostUSD": 0.006,
                },
            }
        ),
        encoding="utf-8",
    )
    (second / "generation_report.json").write_text(
        json.dumps(
            {
                "generatedAt": "2026-06-05T10:00:00Z",
                "sessionId": "session-b",
                "title": "Guide B",
                "model": "claude-sonnet-4-6",
                "pageCount": 6,
                "usage": {
                    "inputTokens": 3000,
                    "outputTokens": 400,
                    "totalTokens": 3400,
                    "estimatedCostUSD": 0.015,
                },
            }
        ),
        encoding="utf-8",
    )
    monkeypatch.setattr(module, "GENERATED_ROOT", generated_root)
    monkeypatch.setattr(module, "USAGE_DB_PATH", tmp_path / "artifacts" / "usage" / "empty.sqlite3")

    daily = module.read_local_usage_summary("day")
    weekly = module.read_local_usage_summary("week")

    assert daily["totals"] == {
        "documents": 2,
        "attempts": 2,
        "failedAttempts": 0,
        "inputTokens": 4000,
        "outputTokens": 600,
        "totalTokens": 4600,
        "pageCount": 10,
        "estimatedCostUSD": 0.021,
        "costPerPageUSD": 0.0021,
    }
    assert [bucket["label"] for bucket in daily["buckets"]] == ["2026-06-04", "2026-06-05"]
    assert daily["buckets"][0]["documents"][0]["sessionId"] == "session-a"
    assert daily["buckets"][0]["documents"][0]["generatedBy"] == {
        "oid": "user-a",
        "name": "Analyst One",
        "username": "analyst.one@example.com",
    }
    assert daily["buckets"][0]["documents"][0]["usage"]["pageCount"] == 4
    assert daily["buckets"][0]["documents"][0]["usage"]["costPerPageUSD"] == 0.0015
    assert weekly["buckets"][0]["label"] == "2026-W23"
    assert weekly["buckets"][0]["totals"]["documents"] == 2
    assert weekly["buckets"][0]["totals"]["pageCount"] == 10


def test_usage_summary_prefers_remote_api_when_configured(monkeypatch: pytest.MonkeyPatch) -> None:
    module = load_app_server()
    monkeypatch.setenv("KCXDOC_REMOTE_API_BASE_URL", "https://kcxdocumentor-ai-dev.azurewebsites.net/")

    def fake_fetch(base_url: str, range_name: str, timeout_seconds: float = 5.0, bearer_token: str = "") -> dict[str, Any]:
        assert base_url == "https://kcxdocumentor-ai-dev.azurewebsites.net"
        assert range_name == "week"
        assert bearer_token == "token-123"
        return {
            "range": "week",
            "generatedAt": "2026-06-05T12:00:00Z",
            "totals": {
                "documents": 1,
                "attempts": 1,
                "failedAttempts": 0,
                "inputTokens": 100,
                "outputTokens": 20,
                "totalTokens": 120,
                "pageCount": 4,
                "estimatedCostUSD": 0.012,
                "costPerPageUSD": 0.003,
            },
            "buckets": [{"label": "2026-W23", "totals": {"documents": 1}, "documents": []}],
        }

    monkeypatch.setattr(module, "fetch_remote_usage_summary", fake_fetch)

    summary = module.read_usage_summary("week", bearer_token="token-123")

    assert summary["range"] == "week"
    assert summary["totals"]["documents"] == 1
    assert summary["totals"]["costPerPageUSD"] == 0.003
    assert summary["buckets"][0]["label"] == "2026-W23"
    assert summary["days"] == []


def test_usage_summary_surfaces_remote_api_failure_when_configured(monkeypatch: pytest.MonkeyPatch) -> None:
    module = load_app_server()
    monkeypatch.setenv("KCXDOC_REMOTE_API_BASE_URL", "https://kcxdocumentor-ai-dev.azurewebsites.net")
    monkeypatch.setattr(module, "fetch_remote_usage_summary", lambda *_args, **_kwargs: (_ for _ in ()).throw(OSError("offline")))

    with pytest.raises(OSError, match="offline"):
        module.read_usage_summary("day", bearer_token="token-123")


def test_usage_summary_backfills_page_count_from_generated_docx(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    module = load_app_server()
    generated_root = tmp_path / "artifacts" / "generated"
    session_dir = generated_root / "session-with-docx"
    session_dir.mkdir(parents=True)
    (session_dir / "generation_report.json").write_text(
        json.dumps(
            {
                "generatedAt": "2026-06-04T16:34:09Z",
                "sessionId": "session-with-docx",
                "title": "Guide With Pages",
                "model": "claude-sonnet-4-6",
                "usage": {
                    "inputTokens": 5000,
                    "outputTokens": 1000,
                    "totalTokens": 6000,
                    "estimatedCostUSD": 0.03,
                },
            }
        ),
        encoding="utf-8",
    )
    doc = Document()
    doc.add_paragraph("First page content.")
    doc.add_page_break()
    doc.add_paragraph("Second page content.")
    doc.add_page_break()
    doc.add_paragraph("Third page content.")
    doc.save(session_dir / "user_guide.anthropic.docx")
    monkeypatch.setattr(module, "GENERATED_ROOT", generated_root)
    monkeypatch.setattr(module, "USAGE_DB_PATH", tmp_path / "artifacts" / "usage" / "empty.sqlite3")

    summary = module.read_local_usage_summary("month")

    assert summary["totals"]["pageCount"] >= 3
    assert summary["totals"]["costPerPageUSD"] > 0
    document = summary["buckets"][0]["documents"][0]
    assert document["usage"]["pageCount"] >= 3
    assert document["usage"]["costPerPageUSD"] > 0


def test_usage_summary_keeps_db_entries_after_session_report_is_deleted(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    module = load_app_server()
    generated_root = tmp_path / "artifacts" / "generated"
    db_path = tmp_path / "artifacts" / "usage" / "generation_usage.sqlite3"
    generated_root.mkdir(parents=True)
    db_path.parent.mkdir(parents=True)
    with sqlite3.connect(db_path) as connection:
        connection.execute(
            """
            CREATE TABLE generation_usage (
                generation_run_id TEXT PRIMARY KEY,
                generated_at TEXT,
                recorded_at TEXT,
                session_id TEXT,
                title TEXT,
                provider TEXT,
                model TEXT,
                prompt_version TEXT,
                input_tokens INTEGER,
                output_tokens INTEGER,
                total_tokens INTEGER,
                estimated_cost_usd REAL,
                report_json TEXT
            )
            """
        )
        connection.execute(
            """
            INSERT INTO generation_usage VALUES (
                'run-a', '2026-06-04T16:34:09Z', '2026-06-04T16:35:00Z',
                'deleted-session', 'Deleted Session Guide', 'anthropic',
                'claude-sonnet-4-6', 'guide-draft-v1', 1200, 300, 1500,
                0.0081, '{}'
            )
            """
        )
    monkeypatch.setattr(module, "GENERATED_ROOT", generated_root)
    monkeypatch.setattr(module, "USAGE_DB_PATH", db_path)

    summary = module.read_local_usage_summary("month")

    assert summary["totals"]["documents"] == 1
    assert summary["totals"]["attempts"] == 1
    assert summary["totals"]["failedAttempts"] == 0
    assert summary["totals"]["totalTokens"] == 1500
    assert summary["totals"]["pageCount"] == 0
    assert summary["totals"]["estimatedCostUSD"] == 0.0081
    assert summary["totals"]["costPerPageUSD"] == 0.0
    assert summary["buckets"][0]["documents"][0]["sessionId"] == "deleted-session"


def test_usage_summary_counts_failed_attempt_spend_without_document(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    module = load_app_server()
    generated_root = tmp_path / "artifacts" / "generated"
    failed_dir = generated_root / "failed-session"
    failed_dir.mkdir(parents=True)
    (failed_dir / "generation_failure.json").write_text(
        json.dumps(
            {
                "schemaVersion": 1,
                "status": "failed",
                "generatedAt": "2026-06-04T16:34:09Z",
                "sessionId": "failed-session",
                "title": "Failed guide generation",
                "model": "claude-sonnet-4-6",
                "provider": "anthropic",
                "promptVersion": "guide-draft-v1",
                "usage": {
                    "inputTokens": 2000,
                    "outputTokens": 500,
                    "totalTokens": 2500,
                    "estimatedCostUSD": 0.0135,
                },
                "errorMessage": "Anthropic returned invalid guide JSON.",
            }
        ),
        encoding="utf-8",
    )
    monkeypatch.setattr(module, "GENERATED_ROOT", generated_root)
    monkeypatch.setattr(module, "USAGE_DB_PATH", tmp_path / "artifacts" / "usage" / "empty.sqlite3")

    summary = module.read_local_usage_summary("day")

    assert summary["totals"]["documents"] == 0
    assert summary["totals"]["attempts"] == 1
    assert summary["totals"]["failedAttempts"] == 1
    assert summary["totals"]["totalTokens"] == 2500
    assert summary["totals"]["pageCount"] == 0
    assert summary["totals"]["estimatedCostUSD"] == 0.0135
    assert summary["totals"]["costPerPageUSD"] == 0.0
    document = summary["buckets"][0]["documents"][0]
    assert document["status"] == "failed"
    assert document["errorMessage"] == "Anthropic returned invalid guide JSON."


def test_export_usage_records_for_remote_uses_sqlite_records(tmp_path: Path) -> None:
    module = load_app_server()
    db_path = tmp_path / "artifacts" / "usage" / "generation_usage.sqlite3"
    db_path.parent.mkdir(parents=True)
    with sqlite3.connect(db_path) as connection:
        connection.execute(
            """
            CREATE TABLE generation_usage (
                generation_run_id TEXT PRIMARY KEY,
                generated_at TEXT,
                recorded_at TEXT,
                session_id TEXT,
                title TEXT,
                provider TEXT,
                model TEXT,
                prompt_version TEXT,
                input_tokens INTEGER,
                output_tokens INTEGER,
                total_tokens INTEGER,
                estimated_cost_usd REAL,
                page_count INTEGER,
                status TEXT,
                error_message TEXT,
                report_json TEXT
            )
            """
        )
        connection.execute(
            """
            INSERT INTO generation_usage VALUES (
                'run-a', '2026-06-04T16:34:09Z', '2026-06-04T16:35:00Z',
                'session-a', 'Guide A', 'anthropic', 'claude-sonnet-4-6',
                'guide-draft-v1', 1000, 200, 1200, 0.006, 4,
                'succeeded', '', '{}'
            )
            """
        )

    payload = module.export_usage_records_for_remote(db_path)

    assert payload["schemaVersion"] == 1
    assert payload["source"] == "local-sqlite"
    assert payload["records"] == [
        {
            "generationRunId": "run-a",
            "generatedAt": "2026-06-04T16:34:09Z",
            "sessionId": "session-a",
            "title": "Guide A",
            "provider": "anthropic",
            "model": "claude-sonnet-4-6",
            "promptVersion": "guide-draft-v1",
            "inputTokens": 1000,
            "outputTokens": 200,
            "totalTokens": 1200,
            "estimatedCostUSD": 0.006,
            "pageCount": 4,
            "status": "succeeded",
            "errorMessage": "",
        }
    ]


def test_migrate_usage_records_posts_local_export_to_remote(monkeypatch: pytest.MonkeyPatch) -> None:
    module = load_app_server()
    posted: dict[str, Any] = {}
    monkeypatch.setenv("KCXDOC_REMOTE_API_BASE_URL", "https://kcxdocumentor-ai-dev.azurewebsites.net")
    monkeypatch.setattr(
        module,
        "export_usage_records_for_remote",
        lambda: {
            "schemaVersion": 1,
            "source": "local-sqlite",
            "records": [{"generationRunId": "run-a", "usage": {"totalTokens": 1200}}],
        },
    )

    def fake_post(base_url: str, records: list[dict[str, Any]], timeout_seconds: float = 15.0, bearer_token: str = "") -> dict[str, Any]:
        posted["base_url"] = base_url
        posted["records"] = records
        posted["bearer_token"] = bearer_token
        posted["timeout_seconds"] = timeout_seconds
        return {"imported": 1, "records": [{"generationRunId": "run-a", "status": "succeeded"}]}

    monkeypatch.setattr(module, "post_remote_usage_records", fake_post)

    result = module.migrate_usage_records(bearer_token="token-123")

    assert result["exported"] == 1
    assert result["imported"] == 1
    assert posted == {
        "base_url": "https://kcxdocumentor-ai-dev.azurewebsites.net",
        "records": [{"generationRunId": "run-a", "usage": {"totalTokens": 1200}}],
        "bearer_token": "token-123",
        "timeout_seconds": 30.0,
    }


def test_write_generation_report_page_count_persists_usage_shape(tmp_path: Path) -> None:
    module = load_app_server()
    report_path = tmp_path / "generation_report.json"
    report_path.write_text(
        json.dumps(
            {
                "generationRunId": "run-a",
                "generatedAt": "2026-06-04T16:34:09Z",
                "sessionId": "session-a",
                "title": "Guide A",
                "usage": {
                    "inputTokens": 1000,
                    "outputTokens": 200,
                    "totalTokens": 1200,
                    "estimatedCostUSD": 0.06,
                },
            }
        ),
        encoding="utf-8",
    )

    updated = module.write_generation_report_page_count(report_path, 12)
    persisted = json.loads(report_path.read_text(encoding="utf-8"))

    assert updated["pageCount"] == 12
    assert updated["usage"]["pageCount"] == 12
    assert updated["usage"]["costPerPageUSD"] == 0.005
    assert persisted == updated


def test_report_page_count_updates_remote_usage(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    module = load_app_server()
    generated_root = tmp_path / "artifacts" / "generated"
    processed_root = tmp_path / "samples" / "processed"
    session_dir = processed_root / "session-a"
    generated_dir = generated_root / "session-a"
    session_dir.mkdir(parents=True)
    generated_dir.mkdir(parents=True)
    (session_dir / "manifest.json").write_text("{}", encoding="utf-8")
    (generated_dir / "generation_report.json").write_text(
        json.dumps(
            {
                "generationRunId": "run-a",
                "generatedAt": "2026-06-04T16:34:09Z",
                "sessionId": "session-a",
                "title": "Guide A",
                "usage": {"estimatedCostUSD": 0.06},
            }
        ),
        encoding="utf-8",
    )
    doc = Document()
    doc.add_paragraph("First page content.")
    doc.add_page_break()
    doc.add_paragraph("Second page content.")
    doc.save(generated_dir / "user_guide.anthropic.docx")
    posted: dict[str, Any] = {}

    monkeypatch.setenv("KCXDOC_REMOTE_API_BASE_URL", "https://kcxdocumentor-ai-dev.azurewebsites.net")
    monkeypatch.setattr(module, "GENERATED_ROOT", generated_root)
    monkeypatch.setattr(module, "PROCESSED_ROOT", processed_root)
    monkeypatch.setattr(module, "USAGE_DB_PATH", tmp_path / "artifacts" / "usage" / "generation_usage.sqlite3")

    def fake_post(base_url: str, records: list[dict[str, Any]], timeout_seconds: float = 15.0, bearer_token: str = "") -> dict[str, Any]:
        posted["base_url"] = base_url
        posted["records"] = records
        posted["bearer_token"] = bearer_token
        return {"imported": 1}

    monkeypatch.setattr(module, "post_remote_usage_records", fake_post)

    result = module.report_page_count({"sessionId": "session-a"}, bearer_token="token-123")

    assert result["pageCount"] >= 2
    assert result["usageUpdate"] == {"imported": 1}
    assert posted["bearer_token"] == "token-123"
    assert posted["records"][0]["pageCount"] == result["pageCount"]
    assert posted["records"][0]["usage"]["pageCount"] == result["pageCount"]


def test_session_source_video_uses_manifest_recording_path(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    module = load_app_server()
    source = tmp_path / "samples" / "raw" / "demo.mp4"
    source.parent.mkdir(parents=True)
    source.write_bytes(b"fake video")

    resolved = module.session_source_path({"sourceFile": str(source)})

    assert resolved == source.resolve()


def test_parse_byte_range_supports_video_scrubbing_requests() -> None:
    module = load_app_server()

    assert module.parse_byte_range("bytes=10-19", 100) == (10, 19)
    assert module.parse_byte_range("bytes=90-", 100) == (90, 99)
    assert module.parse_byte_range("bytes=-10", 100) == (90, 99)
    with pytest.raises(module.HttpError):
        module.parse_byte_range("bytes=100-120", 100)


def test_delete_session_artifacts_removes_processed_and_generated(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    module = load_app_server()
    processed_root = tmp_path / "samples" / "processed"
    generated_root = tmp_path / "artifacts" / "generated"
    session_dir = processed_root / "session-a"
    generated_dir = generated_root / "session-a"
    db_path = tmp_path / "artifacts" / "usage" / "generation_usage.sqlite3"
    session_dir.mkdir(parents=True)
    generated_dir.mkdir(parents=True)
    db_path.parent.mkdir(parents=True)
    (session_dir / "manifest.json").write_text(json.dumps({"schemaVersion": 1, "sessionId": "session-a"}), encoding="utf-8")
    (generated_dir / "user_guide.anthropic.docx").write_bytes(b"docx")
    db_path.write_bytes(b"sqlite placeholder")
    monkeypatch.setattr(module, "PROCESSED_ROOT", processed_root)
    monkeypatch.setattr(module, "GENERATED_ROOT", generated_root)
    monkeypatch.setattr(module, "USAGE_DB_PATH", db_path)

    result = module.delete_session_artifacts({"sessionId": "session-a"})

    assert result["sessionId"] == "session-a"
    assert len(result["deleted"]) == 2
    assert any(path.endswith("samples/processed/session-a") for path in result["deleted"])
    assert any(path.endswith("artifacts/generated/session-a") for path in result["deleted"])
    assert not session_dir.exists()
    assert not generated_dir.exists()
    assert db_path.exists()
    assert db_path.read_bytes() == b"sqlite placeholder"
    assert result["sessions"] == []


def test_delete_session_artifacts_rejects_path_escape(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    module = load_app_server()
    processed_root = tmp_path / "samples" / "processed"
    generated_root = tmp_path / "artifacts" / "generated"
    outside_dir = tmp_path / "outside-session"
    protected_file = outside_dir / "keep.txt"
    processed_root.mkdir(parents=True)
    generated_root.mkdir(parents=True)
    outside_dir.mkdir()
    protected_file.write_text("do not delete", encoding="utf-8")
    try:
        (processed_root / "session-a").symlink_to(outside_dir, target_is_directory=True)
    except OSError:
        pytest.skip("filesystem does not allow directory symlinks")
    monkeypatch.setattr(module, "PROCESSED_ROOT", processed_root)
    monkeypatch.setattr(module, "GENERATED_ROOT", generated_root)

    with pytest.raises(module.HttpError) as exc:
        module.delete_session_artifacts({"sessionId": "session-a"})

    assert exc.value.status == 400
    assert protected_file.exists()


def test_processing_command_builder_uses_python_and_no_media_tools_when_requested(tmp_path: Path) -> None:
    module = load_app_server()
    builder = find_callable(module, "build_process_command", "build_processing_command")
    recording = tmp_path / "samples" / "raw" / "example.mp4"
    recording.parent.mkdir(parents=True)
    recording.write_bytes(b"placeholder")

    command = invoke_with_supported_kwargs(
        builder,
        recording,
        repo_root=ROOT,
        output_root=tmp_path / "samples" / "processed",
        session_id="demo-session",
        target_application="Enterprise Rx",
        no_media_tools=True,
    )

    command_text = " ".join(str(part) for part in command)
    assert str(APP_SERVER.parent / "process_recording.py") in command_text
    assert "--session-id" in command
    assert "demo-session" in command
    assert "--target-application" in command
    assert "Enterprise Rx" in command
    assert "--no-media-tools" in command
    assert command[0] == sys.executable or command[0].endswith("python") or command[0].endswith("python3")


def test_recording_import_helper_copies_to_raw_root_without_path_escape(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    module = load_app_server()
    importer = find_optional_callable(module, "import_recording", "copy_imported_recording", "save_uploaded_recording")
    if importer is None:
        pytest.skip("app server does not expose a recording import/upload helper yet")

    raw_root = tmp_path / "samples" / "raw"
    raw_root.mkdir(parents=True)
    source = tmp_path / "source" / "training clip.mp4"
    source.parent.mkdir()
    source.write_bytes(b"fake video bytes")
    monkeypatch.setattr(module, "RAW_ROOT", raw_root)

    result = invoke_with_supported_kwargs(importer, source, raw_root=raw_root, filename="training clip.mp4")
    imported_path = Path(result["path"] if isinstance(result, dict) and "path" in result else result)

    assert imported_path.resolve().parent == raw_root.resolve()
    assert imported_path.name.endswith(".mp4")
    assert imported_path.read_bytes() == b"fake video bytes"

    outside = tmp_path / "outside.mp4"
    with pytest.raises((PermissionError, ValueError, module.HttpError if hasattr(module, "HttpError") else Exception)):
        invoke_with_supported_kwargs(importer, source, raw_root=raw_root, filename="../outside.mp4")
    assert not outside.exists()


def test_transcript_command_builder_includes_sidecar_when_helper_supports_it(tmp_path: Path) -> None:
    module = load_app_server()
    builder = find_callable(module, "build_process_command", "build_processing_command")
    if "transcript" not in inspect.signature(builder).parameters and "transcript_path" not in inspect.signature(builder).parameters:
        pytest.skip("process command builder does not expose transcript sidecar support yet")

    recording = tmp_path / "samples" / "raw" / "example.mp4"
    transcript = tmp_path / "samples" / "raw" / "example.txt"
    recording.parent.mkdir(parents=True)
    recording.write_bytes(b"placeholder")
    transcript.write_text("Click Save to finish the workflow.", encoding="utf-8")

    command = invoke_with_supported_kwargs(
        builder,
        recording,
        repo_root=ROOT,
        output_root=tmp_path / "samples" / "processed",
        session_id="demo-session",
        target_application="Enterprise Rx",
        no_media_tools=True,
        transcript=transcript,
        transcript_path=transcript,
    )

    command_text = " ".join(str(part) for part in command)
    assert "--transcript" in command
    assert str(transcript) in command_text


def test_processing_command_builder_can_enable_teams_recording_profile(tmp_path: Path) -> None:
    module = load_app_server()
    builder = find_callable(module, "build_process_command", "build_processing_command")
    recording = tmp_path / "samples" / "raw" / "teams.mp4"
    recording.parent.mkdir(parents=True)
    recording.write_bytes(b"placeholder")

    command = invoke_with_supported_kwargs(
        builder,
        recording,
        repo_root=ROOT,
        output_root=tmp_path / "samples" / "processed",
        target_application="SendKey",
        source_profile="teams-recording",
    )

    assert "--source-profile" in command
    assert "teams-recording" in command


def test_health_or_tooling_helper_reports_expected_local_dependencies() -> None:
    module = load_app_server()
    checker = find_optional_callable(module, "get_health", "health_check", "check_tooling", "get_tooling_status")
    if checker is None:
        pytest.skip("app server does not expose a health/tooling helper yet")

    status = checker()

    assert isinstance(status, dict)
    serialized = json.dumps(status).lower()
    assert "ffmpeg" in serialized
    assert "ffprobe" in serialized
    assert "whisper" in serialized
    assert "modelavailable" in serialized
    if "anthropic" in serialized or "api" in serialized:
        assert "sk-ant-" not in serialized
    assert "ok" in status or "ready" in status or "tools" in status or "dependencies" in status


def test_health_reports_configured_external_whisper_share(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    module = load_app_server()
    whisper_cli = tmp_path / "external" / "whisper" / "bin" / "whisper-cli"
    whisper_model = tmp_path / "external" / "whisper" / "models" / "ggml-base.en.bin"
    whisper_cli.parent.mkdir(parents=True)
    whisper_model.parent.mkdir(parents=True)
    whisper_cli.write_text("#!/bin/sh\n", encoding="utf-8")
    whisper_model.write_bytes(b"model")
    monkeypatch.setenv("KCXDOC_WHISPER_CLI", str(whisper_cli))
    monkeypatch.setenv("KCXDOC_WHISPER_MODEL", str(whisper_model))

    status = module.get_health()

    assert status["tools"]["whisper"]["available"] is True
    assert status["tools"]["whisper"]["path"] == str(whisper_cli)
    assert status["tools"]["whisper"]["modelAvailable"] is True
    assert status["tools"]["whisper"]["modelPath"] == str(whisper_model)


def test_frame_review_actions_persist_and_merge_into_session(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    module = load_app_server()
    processed_root = tmp_path / "samples" / "processed"
    session_dir = processed_root / "review-session"
    session_dir.mkdir(parents=True)
    (session_dir / "manifest.json").write_text(json.dumps({"schemaVersion": 1, "sessionId": "review-session"}), encoding="utf-8")
    (session_dir / "frame_scores.json").write_text(
        json.dumps(
            {
                "schemaVersion": 1,
                "frames": [
                    {
                        "id": "frame-0001",
                        "timestampSeconds": 12.0,
                        "timestamp": "00:12",
                        "path": "frames/candidates/frame-0001.png",
                        "webPath": "frames/candidates/frame-0001.png",
                        "created": True,
                        "score": 0.8,
                    }
                ],
            }
        ),
        encoding="utf-8",
    )
    (session_dir / "procedure_trace.json").write_text(
        json.dumps(
            {
                "schemaVersion": 1,
                "segments": [
                    {
                        "id": "seg-0001",
                        "candidateImages": [
                            {
                                "frameId": "frame-0001",
                                "reviewStatus": "pending",
                            }
                        ],
                    }
                ],
            }
        ),
        encoding="utf-8",
    )
    monkeypatch.setattr(module, "PROCESSED_ROOT", processed_root)
    monkeypatch.setattr(module, "GENERATED_ROOT", tmp_path / "artifacts" / "generated")

    result = module.update_frame_review(
        {
            "sessionId": "review-session",
            "frameId": "frame-0001",
            "action": "approve",
            "note": "Good application state.",
            "assignedSegmentId": "seg-0001",
        }
    )

    assert result["frameReview"]["summary"]["approved"] == 1
    review_file = json.loads((session_dir / "frame_review.json").read_text(encoding="utf-8"))
    assert review_file["frames"]["frame-0001"]["status"] == "approved"
    assert review_file["frames"]["frame-0001"]["note"] == "Good application state."
    assert review_file["frames"]["frame-0001"]["assignedSegmentId"] == "seg-0001"

    session = module.read_session(session_dir)
    image = session["procedureTrace"]["segments"][0]["candidateImages"][0]
    assert image["reviewStatus"] == "approved"
    assert image["reviewNote"] == "Good application state."
    assert session["frameReview"]["frames"][0]["assignedSegmentId"] == "seg-0001"


def test_frame_review_accepts_ui_status_and_review_note_aliases(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    module = load_app_server()
    processed_root = tmp_path / "samples" / "processed"
    session_dir = processed_root / "review-session"
    session_dir.mkdir(parents=True)
    (session_dir / "frame_scores.json").write_text(
        json.dumps(
            {
                "schemaVersion": 1,
                "frames": [
                    {
                        "id": "frame-0001",
                        "timestampSeconds": 12.0,
                        "path": "frames/candidates/frame-0001.png",
                        "created": True,
                    }
                ],
            }
        ),
        encoding="utf-8",
    )
    monkeypatch.setattr(module, "PROCESSED_ROOT", processed_root)

    module.update_frame_review(
        {
            "sessionId": "review-session",
            "frameId": "frame-0001",
            "reviewStatus": "approved",
            "reviewNote": "Use this screenshot in the guide.",
            "assignedSegmentId": "seg-0001",
        }
    )

    review_file = json.loads((session_dir / "frame_review.json").read_text(encoding="utf-8"))
    assert review_file["frames"]["frame-0001"]["status"] == "approved"
    assert review_file["frames"]["frame-0001"]["note"] == "Use this screenshot in the guide."
    assert review_file["frames"]["frame-0001"]["assignedSegmentId"] == "seg-0001"


def test_frame_review_rejects_unknown_or_unsafe_frame_ids(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    module = load_app_server()
    processed_root = tmp_path / "samples" / "processed"
    session_dir = processed_root / "review-session"
    session_dir.mkdir(parents=True)
    (session_dir / "frame_scores.json").write_text(json.dumps({"schemaVersion": 1, "frames": []}), encoding="utf-8")
    monkeypatch.setattr(module, "PROCESSED_ROOT", processed_root)

    with pytest.raises(module.HttpError):
        module.update_frame_review({"sessionId": "review-session", "frameId": "../frame", "action": "approve"})
    with pytest.raises(module.HttpError):
        module.update_frame_review({"sessionId": "review-session", "frameId": "frame-9999", "action": "approve"})


def test_extract_review_frame_adds_candidate_and_review_entry(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    module = load_app_server()
    processed_root = tmp_path / "samples" / "processed"
    session_dir = processed_root / "extract-session"
    session_dir.mkdir(parents=True)
    source = tmp_path / "samples" / "raw" / "demo.mov"
    source.parent.mkdir(parents=True)
    source.write_bytes(b"fake video")
    (session_dir / "manifest.json").write_text(
        json.dumps(
            {
                "schemaVersion": 1,
                "sessionId": "extract-session",
                "sourceFile": str(source),
                "processing": {"frameCropFilter": "crop=iw:ih:0:0"},
            }
        ),
        encoding="utf-8",
    )
    (session_dir / "frame_scores.json").write_text(json.dumps({"schemaVersion": 1, "frames": []}), encoding="utf-8")
    monkeypatch.setattr(module, "PROCESSED_ROOT", processed_root)
    monkeypatch.setattr(module.shutil, "which", lambda name: "/usr/bin/ffmpeg" if name == "ffmpeg" else None)

    def fake_run_command(command: list[str]) -> dict[str, Any]:
        Path(command[-1]).write_bytes(b"png")
        return {"returnCode": 0, "stdout": "", "stderr": ""}

    monkeypatch.setattr(module, "run_command", fake_run_command)

    result = module.extract_review_frame(
        {
            "sessionId": "extract-session",
            "timestamp": "01:02",
            "frameId": "review-frame-0007",
            "reviewStatus": "approved",
            "reviewNote": "Use this dialog.",
            "segmentId": "seg-0002",
        }
    )

    assert result["frame"]["id"] == "review-frame-0007"
    assert result["frame"]["timestampSeconds"] == 62.0
    assert (session_dir / "frames" / "candidates" / "review-frame-0007.png").read_bytes() == b"png"
    frame_scores = json.loads((session_dir / "frame_scores.json").read_text(encoding="utf-8"))
    assert frame_scores["frames"][0]["source"] == "manual-review-extract"
    assert frame_scores["frames"][0]["cropFilter"] == "crop=iw:ih:0:0"
    review = json.loads((session_dir / "frame_review.json").read_text(encoding="utf-8"))
    assert review["frames"]["review-frame-0007"]["status"] == "approved"
    assert review["frames"]["review-frame-0007"]["note"] == "Use this dialog."
    assert review["frames"]["review-frame-0007"]["assignedSegmentId"] == "seg-0002"


def test_extract_review_frame_runs_ocr_and_saves_prompt_context(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    module = load_app_server()
    processed_root = tmp_path / "samples" / "processed"
    session_dir = processed_root / "extract-context"
    session_dir.mkdir(parents=True)
    source = tmp_path / "samples" / "raw" / "demo.mov"
    source.parent.mkdir(parents=True)
    source.write_bytes(b"fake video")
    (session_dir / "manifest.json").write_text(
        json.dumps({"schemaVersion": 1, "sessionId": "extract-context", "sourceFile": str(source)}),
        encoding="utf-8",
    )
    (session_dir / "procedure_trace.json").write_text(
        json.dumps(
            {
                "schemaVersion": 1,
                "sessionId": "extract-context",
                "recording": {"targetApplication": "Blink Rx"},
                "segments": [
                    {
                        "id": "seg-0002",
                        "speakerText": "Click Submit to finish the refill request.",
                        "startSeconds": 60.0,
                        "endSeconds": 90.0,
                        "start": "00:01:00.000",
                        "end": "00:01:30.000",
                        "candidateImages": [],
                    }
                ],
            }
        ),
        encoding="utf-8",
    )
    (session_dir / "frame_scores.json").write_text(json.dumps({"schemaVersion": 1, "frames": []}), encoding="utf-8")
    monkeypatch.setattr(module, "PROCESSED_ROOT", processed_root)
    monkeypatch.setattr(module.shutil, "which", lambda name: f"/usr/bin/{name}" if name in {"ffmpeg", "tesseract"} else None)

    def fake_run_command(command: list[str]) -> dict[str, Any]:
        Path(command[-1]).write_bytes(b"png")
        return {"returnCode": 0, "stdout": "", "stderr": ""}

    class FakeProcessor:
        @staticmethod
        def evaluate_frame_visual_quality(frame: dict[str, Any], session_dir: Path) -> dict[str, Any]:
            return {"qualityScore": 0.9, "blurState": "sharp", "averageHash": "ffff0000ffff0000"}

        @staticmethod
        def nearest_visual_duplicate(current_hash: str, prior_hashes: list[tuple[str, str]]) -> None:
            return None

        @staticmethod
        def frame_selection_reason(visual: dict[str, Any], duplicate_of: str | None) -> str:
            return "Added with local visual quality scoring."

        @staticmethod
        def clamp01(value: float) -> float:
            return round(min(1.0, max(0.0, value)), 3)

        @staticmethod
        def parse_float(value: Any) -> float | None:
            return float(value) if value not in (None, "") else None

        @staticmethod
        def run_tesseract_ocr(*args: Any, **kwargs: Any) -> dict[str, Any]:
            return {"textBlocks": [{"text": "Submit Refill Request", "confidence": 88.0}], "error": None}

        @staticmethod
        def average_block_confidence(blocks: list[dict[str, Any]]) -> float:
            return 0.88

        @staticmethod
        def build_placeholder_ocr_frame(frame: dict[str, Any], target_application: str, reason: str) -> dict[str, Any]:
            return {"frameId": frame["id"], "source": "prototype-placeholder", "confidence": 0, "combinedText": target_application, "textBlocks": [], "error": reason}

        @staticmethod
        def build_candidate_image(frame: dict[str, Any], ocr_frame: dict[str, Any], segment: dict[str, Any], target_application: str) -> dict[str, Any]:
            return {
                "frameId": frame["id"],
                "frameEvidenceScore": 0.91,
                "ocrConfidence": 0.88,
                "ocrRelevanceScore": 0.8,
                "ocrClass": "application",
                "ocrText": ocr_frame["combinedText"],
                "contentType": "application",
                "reason": "OCR matched the segment context.",
            }

        @staticmethod
        def assign_frame_recommendation_groups(images: list[dict[str, Any]]) -> list[dict[str, Any]]:
            images[0]["recommendationGroup"] = "recommended"
            images[0]["selectionDecision"] = "recommended"
            images[0]["recommendationReason"] = "Best manual screenshot candidate."
            return images

    monkeypatch.setattr(module, "load_process_recording_module", lambda: FakeProcessor)
    monkeypatch.setattr(module, "run_command", fake_run_command)

    result = module.extract_review_frame(
        {
            "sessionId": "extract-context",
            "timestamp": "01:10",
            "frameId": "review-frame-0008",
            "segmentId": "seg-0002",
        }
    )

    assert result["frame"]["ocrText"] == "Submit Refill Request"
    assert result["frame"]["ocrClass"] == "application"
    assert result["frame"]["recommendationGroup"] == "recommended"
    frame_scores = json.loads((session_dir / "frame_scores.json").read_text(encoding="utf-8"))
    assert frame_scores["frames"][0]["frameEvidenceScore"] == 0.91
    ocr = json.loads((session_dir / "ocr.json").read_text(encoding="utf-8"))
    assert ocr["frames"][0]["combinedText"] == "Submit Refill Request"
    trace = module.apply_frame_review_to_trace(
        json.loads((session_dir / "procedure_trace.json").read_text(encoding="utf-8")),
        json.loads((session_dir / "frame_review.json").read_text(encoding="utf-8")),
    )
    image = trace["segments"][0]["candidateImages"][0]
    assert image["frameId"] == "review-frame-0008"
    assert image["ocrText"] == "Submit Refill Request"
    assert image["recommendationGroup"] == "recommended"
