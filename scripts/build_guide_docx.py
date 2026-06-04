#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path
from typing import Any


WORKSPACE = Path(__file__).resolve().parents[1]
if str(WORKSPACE) not in sys.path:
    sys.path.insert(0, str(WORKSPACE))

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
    try:
        from PIL import Image, ImageStat
    except ImportError:  # pragma: no cover - optional visual QA dependency
        Image = None
        ImageStat = None
    from tools.document_lib.keycentrix_docx import (
        FONT,
        add_cover,
        add_header_footer,
        add_metadata_table,
        style_document,
    )
except ImportError as exc:
    raise SystemExit(
        "Missing DOCX dependency. Install python-docx in the active Python environment, "
        "then rerun this command. Example: python3 -m pip install python-docx"
    ) from exc


@dataclass
class GuideStep:
    title: str
    action: str = ""
    expected_result: str = ""
    section_title: str = ""
    notes: list[str] = field(default_factory=list)
    ui_text: list[str] = field(default_factory=list)
    action_hints: list[str] = field(default_factory=list)
    transcript: str = ""
    start: str = ""
    end: str = ""
    screenshot: Path | None = None
    screenshot_caption: str = ""
    reviewer_comments: list[str] = field(default_factory=list)


@dataclass
class GuideDraft:
    title: str
    version: str
    status: str
    owner: str
    effective_date: str
    summary: str
    audience: list[str]
    prerequisites: list[str]
    workflow_overview: list[str]
    expected_results: list[str]
    troubleshooting: list[str]
    review_notes: list[str]
    source_metadata: dict[str, str]
    generation_metadata: dict[str, str]
    steps: list[GuideStep]


def as_list(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, list):
        return [stringify_list_item(item) for item in value if stringify_list_item(item)]
    if isinstance(value, tuple):
        return [stringify_list_item(item) for item in value if stringify_list_item(item)]
    text = str(value).strip()
    return [text] if text else []


def stringify_list_item(item: Any) -> str:
    if isinstance(item, dict):
        parts = []
        for key in (
            "id",
            "severity",
            "category",
            "description",
            "resolution",
            "message",
            "title",
            "segmentId",
            "segmentIds",
            "sourceSegment",
            "sourceSegments",
            "timestamp",
            "timestampSeconds",
            "start",
            "end",
            "startSeconds",
            "endSeconds",
            "reason",
            "reasons",
            "note",
            "notes",
            "reviewNote",
            "reviewNotes",
            "reviewGuidance",
            "reviewerGuidance",
        ):
            value = item.get(key)
            if value is not None and str(value).strip():
                label = key.replace("_", " ").title()
                if isinstance(value, (dict, list, tuple)):
                    rendered = json.dumps(value, ensure_ascii=False)
                else:
                    rendered = str(value).strip()
                parts.append(f"{label}: {rendered}")
        if parts:
            return " | ".join(parts)
    return str(item).strip()


def pick_value(source: dict[str, Any], *keys: str) -> Any:
    for key in keys:
        value = source.get(key)
        if value is not None:
            return value
    return None


def text_value(source: dict[str, Any], *keys: str, default: str = "") -> str:
    for key in keys:
        value = source.get(key)
        if value is not None and str(value).strip():
            return str(value).strip()
    return default


def first_text_value(*sources: tuple[dict[str, Any], tuple[str, ...]], default: str = "") -> str:
    for source, keys in sources:
        value = text_value(source, *keys)
        if value:
            return value
    return default


def is_internal_pipeline_text(value: str) -> bool:
    text = value.strip().lower()
    if not text:
        return True
    blocked_phrases = (
        "generated from a local procedure trace",
        "local procedure trace",
        "source recording context before publishing",
        "before publishing",
        "application reviewer",
        "extracted procedure steps",
        "original recording",
    )
    return any(phrase in text for phrase in blocked_phrases)


def format_duration(value: Any) -> str:
    if value is None or not str(value).strip():
        return "Not specified"
    if isinstance(value, str):
        stripped = value.strip()
        if re.fullmatch(r"\d{1,2}:\d{2}(?::\d{2})?", stripped):
            return stripped
        try:
            seconds = float(stripped)
        except ValueError:
            return stripped
    else:
        try:
            seconds = float(value)
        except (TypeError, ValueError):
            return str(value).strip()
    if seconds < 0:
        return "Not specified"
    rounded = int(round(seconds))
    hours, remainder = divmod(rounded, 3600)
    minutes, secs = divmod(remainder, 60)
    return f"{hours:02d}:{minutes:02d}:{secs:02d}"


def format_int(value: Any) -> str:
    if value is None or not str(value).strip():
        return "Not specified"
    try:
        return f"{int(float(str(value).replace(',', '').strip())):,}"
    except ValueError:
        return str(value).strip()


def format_cost(value: Any) -> str:
    if value is None or not str(value).strip():
        return "Not specified"
    text = str(value).strip()
    if text.startswith("$"):
        return text
    try:
        return f"${float(text):.4f}"
    except ValueError:
        return text


def model_label(value: Any) -> str:
    if isinstance(value, dict):
        return text_value(value, "id", "name", "model", "modelId", "provider", default="Not specified")
    if value is None or not str(value).strip():
        return "Not specified"
    return str(value).strip()


def usage_value(usage: dict[str, Any], *keys: str) -> Any:
    for key in keys:
        value = usage.get(key)
        if value is not None:
            return value
    return None


def generation_metadata(data: dict[str, Any]) -> dict[str, str]:
    usage = data.get("usage") if isinstance(data.get("usage"), dict) else {}
    generation = data.get("generation") if isinstance(data.get("generation"), dict) else {}
    model = pick_value(data, "model", "modelId", "modelName") or pick_value(generation, "model", "modelId", "modelName")
    input_tokens = usage_value(usage, "inputTokens", "input_tokens", "cacheReadInputTokens", "cache_read_input_tokens")
    output_tokens = usage_value(usage, "outputTokens", "output_tokens")
    total_tokens = usage_value(usage, "totalTokens", "total_tokens")
    if total_tokens is None:
        try:
            total_tokens = (
                int(float(str(input_tokens or 0))) + int(float(str(output_tokens or 0)))
                if input_tokens is not None or output_tokens is not None
                else None
            )
        except ValueError:
            total_tokens = None
    estimated_cost = (
        usage_value(usage, "estimatedCostUSD", "estimated_cost_usd", "costUSD", "cost_usd")
        or pick_value(generation, "estimatedCostUSD", "estimated_cost_usd", "costUSD", "cost_usd")
    )
    generated_at = (
        text_value(data, "generatedAt", "generated_at", "createdUtc", "createdAt")
        or text_value(generation, "generatedAt", "generated_at", "createdUtc", "createdAt")
    )

    metadata = {
        "Model": model_label(model),
        "Generated At": generated_at or "Not specified",
        "Input Tokens": format_int(input_tokens),
        "Output Tokens": format_int(output_tokens),
        "Total Tokens": format_int(total_tokens),
        "Estimated Cost": format_cost(estimated_cost),
    }
    return {key: value for key, value in metadata.items() if value != "Not specified"}


def split_markdown_bold(text: str) -> list[tuple[str, bool]]:
    parts: list[tuple[str, bool]] = []
    position = 0
    bold = False
    for match in re.finditer(r"\*\*", text):
        if match.start() > position:
            parts.append((text[position:match.start()], bold))
        bold = not bold
        position = match.end()
    if position < len(text):
        parts.append((text[position:], bold))
    if bold:
        return [(text, False)]
    return [(part, is_bold) for part, is_bold in parts if part]


def add_rich_text(paragraph: Any, text: str, *, bold: bool = False, italic: bool = False, size: Pt | None = None) -> None:
    for part, part_bold in split_markdown_bold(text):
        run = paragraph.add_run(part)
        run.font.name = FONT
        if size is not None:
            run.font.size = size
        run.bold = bold or part_bold
        run.italic = italic


def add_rich_paragraph(doc: Document, text: str, style: str | None = None) -> Any:
    paragraph = doc.add_paragraph(style=style)
    add_rich_text(paragraph, text)
    return paragraph


def add_rich_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        add_rich_paragraph(doc, item, style="List Bullet")


def iter_document_paragraphs(doc: Document) -> list[Any]:
    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    for section in doc.sections:
        paragraphs.extend(section.header.paragraphs)
        paragraphs.extend(section.footer.paragraphs)
    return paragraphs


def apply_markdown_bold_to_paragraph(paragraph: Any) -> None:
    text = paragraph.text
    if "**" not in text:
        return
    parts = split_markdown_bold(text)
    if len(parts) == 1 and parts[0][0] == text and not parts[0][1]:
        return
    base_run = paragraph.runs[0] if paragraph.runs else None
    base = {
        "bold": bool(base_run.bold) if base_run is not None else False,
        "italic": bool(base_run.italic) if base_run is not None else False,
        "font_name": base_run.font.name if base_run is not None else FONT,
        "font_size": base_run.font.size if base_run is not None else None,
        "font_color": base_run.font.color.rgb if base_run is not None else None,
    }
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    for part, part_bold in parts:
        run = paragraph.add_run(part)
        run.bold = base["bold"] or part_bold
        run.italic = base["italic"]
        run.font.name = base["font_name"] or FONT
        if base["font_size"] is not None:
            run.font.size = base["font_size"]
        if base["font_color"] is not None:
            run.font.color.rgb = base["font_color"]


def apply_markdown_bold(doc: Document) -> None:
    for paragraph in iter_document_paragraphs(doc):
        apply_markdown_bold_to_paragraph(paragraph)


def customer_safe_summary(
    data: dict[str, Any],
    document: dict[str, Any],
    recording: dict[str, Any],
    source_recording: dict[str, Any],
    session: dict[str, Any],
    meta: dict[str, Any],
    metadata: dict[str, Any],
    title: str,
) -> str:
    introduction = data.get("introduction") if isinstance(data.get("introduction"), dict) else {}
    candidates = [
        text_value(document, "description", "summary", "purpose"),
        text_value(introduction, "text", "description", "summary"),
        text_value(data, "description", "summary", "purpose"),
    ]
    for candidate in candidates:
        if candidate and not is_internal_pipeline_text(candidate):
            return candidate

    app_name = first_text_value(
        (document, ("targetApplication", "application", "appName", "app_name")),
        (source_recording, ("targetApplication", "application", "appName", "app_name")),
        (recording, ("targetApplication", "application", "appName", "app_name")),
        (data, ("targetApplication", "application", "appName", "app_name")),
        (meta, ("targetApplication", "application", "appName", "app_name")),
        (session, ("targetApplication", "application", "appName", "app_name")),
        (metadata, ("targetApplication", "application", "appName", "app_name")),
    )
    workflow_title = re.sub(r"\b(user|workflow)?\s*guide\b", "", title, flags=re.IGNORECASE).strip(" -:") or "the documented workflow"
    if app_name and app_name.lower() not in workflow_title.lower():
        return f"This guide covers {workflow_title} in {app_name}."
    return f"This guide covers {workflow_title}."


def clean_prerequisites(items: list[str], app_name: str) -> list[str]:
    cleaned = []
    for item in items:
        text = re.sub(r"\s+", " ", item).strip()
        if not text or is_internal_pipeline_text(text):
            continue
        cleaned.append(text)
    if cleaned:
        return cleaned
    app_label = app_name if app_name and app_name != "Not specified" else "the target application"
    return [f"You have access to {app_label} and the permissions needed to complete this workflow."]


def clean_customer_items(items: list[str]) -> list[str]:
    return [item for item in items if item and not is_internal_pipeline_text(item)]


def ambiguous_term_comments(text: str) -> list[str]:
    comments: list[str] = []
    for term in ("PV1", "PDR"):
        if re.search(rf"\b{re.escape(term)}\b", text):
            comments.append(f"Verify that '{term}' is the correct UI term and define it for the intended audience if needed.")
    if re.search(r"\b(?:all\s+)?required checkboxes\b", text, flags=re.IGNORECASE):
        comments.append("Verify the exact required checkbox labels and update the instruction with the confirmed UI names.")
    return comments


def clean_step_title(raw_title: str, index: int, action: str = "") -> str:
    title = re.sub(r"\s+", " ", raw_title or "").strip()
    title = re.sub(r"^\s*Step\s+\d+\s*(?:[-:–—]\s*)?", "", title, flags=re.IGNORECASE)
    title = re.sub(r"\s*:\s*Step\s+\d+\s*$", "", title, flags=re.IGNORECASE)
    title = re.sub(r"\s*[-:–—]\s*Step\s+\d+\s*$", "", title, flags=re.IGNORECASE)
    title = re.sub(r"\s*\(\s*Step\s+\d+\s*\)\s*$", "", title, flags=re.IGNORECASE).strip()
    if not title or re.fullmatch(r"(?:step\s*)?\d+", title, flags=re.IGNORECASE):
        title = derive_title_from_action(action, index)
    return title


def derive_title_from_action(action: str, index: int) -> str:
    words = re.findall(r"[A-Za-z0-9']+", action or "")
    if not words:
        return f"Complete Procedure Step {index}"
    stop_words = {"the", "a", "an", "to", "and", "or", "then", "in", "on", "for", "of", "with"}
    title_words: list[str] = []
    for word in words:
        title_words.append(word.capitalize() if word.lower() not in stop_words else word.lower())
        if len(title_words) >= 7:
            break
    return " ".join(title_words).strip() or f"Complete Procedure Step {index}"


def extend_comment_list(comments: list[str], value: Any) -> None:
    for item in as_list(value):
        if item not in comments:
            comments.append(item)


def resolve_asset_path(raw_path: Any, input_path: Path, asset_roots: list[Path] | None = None) -> Path | None:
    if raw_path is None:
        return None
    if isinstance(raw_path, dict):
        raw_path = raw_path.get("path") or raw_path.get("relativePath") or raw_path.get("file") or raw_path.get("filename")
    if not str(raw_path).strip():
        return None

    candidate = Path(str(raw_path))
    if candidate.is_absolute() and candidate.exists():
        return candidate

    search_roots = [input_path.parent, *(asset_roots or []), WORKSPACE, Path.cwd()]
    for root in search_roots:
        resolved = (root / candidate).resolve()
        if resolved.exists():
            return resolved
    return None


def first_existing_image(step: dict[str, Any], input_path: Path, asset_roots: list[Path] | None = None) -> tuple[Path | None, str]:
    image_values: list[Any] = []
    for key in ("screenshot", "screenshotRef", "image", "selected_image", "selected_screenshot", "selectedImage", "selectedScreenshot"):
        if step.get(key):
            image_values.append(step[key])
    candidate_images = pick_value(step, "candidate_images", "candidateImages")
    if isinstance(candidate_images, list):
        image_values.extend(candidate_images)
    elif candidate_images:
        image_values.append(candidate_images)
    screenshots = step.get("screenshots")
    if isinstance(screenshots, list):
        image_values.extend(screenshots)
    elif screenshots:
        image_values.append(screenshots)

    for value in image_values:
        resolved = resolve_asset_path(value, input_path, asset_roots)
        if resolved:
            caption = text_value(step, "screenshot_caption", "caption", default=resolved.name)
            return resolved, caption
    return None, ""


def is_likely_non_application_screenshot(path: Path | None) -> bool:
    if not path or Image is None or ImageStat is None:
        return False
    try:
        with Image.open(path) as image:
            sample = image.convert("L").resize((32, 32))
            histogram = sample.histogram()
            total_pixels = sum(histogram) or 1
            dark_ratio = sum(histogram[:32]) / total_pixels
            mean_luminance = ImageStat.Stat(sample).mean[0]
    except Exception:
        return False
    return dark_ratio >= 0.75 and mean_luminance < 45


def normalize_step(step: dict[str, Any], index: int, input_path: Path, asset_roots: list[Path] | None = None) -> GuideStep:
    screenshot, caption = first_existing_image(step, input_path, asset_roots)
    transcript = text_value(step, "speaker_text", "speakerText", "transcript", "narration", "description")
    action = text_value(step, "action", "instruction", "intent", "stepTextPlaceholder", "body", default=transcript)
    title = clean_step_title(text_value(step, "title", "name", "shellId", default=f"Step {index}"), index, action)
    expected_result = text_value(step, "expected_result", "expectedResult", "result", "outcome")
    notes = as_list(step.get("notes") or step.get("warnings") or step.get("reviewNotes"))
    ui_text = as_list(pick_value(step, "visible_ui_text", "visibleUiText", "ui_text", "uiText", "ocr_text", "ocrText", "confirmedUiLabels"))
    action_hints = as_list(pick_value(step, "action_hints", "actionHints", "events", "actionHint", "actionHints"))
    reviewer_comments = list(notes)
    if step.get("needsHumanReview") is True:
        reviewer_comments.append("This step was flagged for human review by the draft generator.")
    extend_comment_list(reviewer_comments, step.get("reviewGuidance"))
    extend_comment_list(reviewer_comments, step.get("reviewerGuidance"))
    extend_comment_list(reviewer_comments, step.get("reviewComments"))
    extend_comment_list(reviewer_comments, step.get("humanReview"))
    extend_comment_list(reviewer_comments, step.get("openReviewItems"))
    confidence = step.get("confidence")
    if isinstance(confidence, dict):
        reasons = as_list(confidence.get("reasons"))
        if confidence.get("needsHumanReview") is True:
            reviewer_comments.append("Source confidence requires human review.")
        reviewer_comments.extend(reasons)
    if step.get("reviewStatus") == "rejected":
        reviewer_comments.append("The selected screenshot was rejected during frame review and should be replaced.")
    if ui_text:
        reviewer_comments.append(f"Source UI evidence: {'; '.join(ui_text)}")
    if action_hints:
        reviewer_comments.append(f"Source action hints: {'; '.join(action_hints)}")
    reviewer_comments.extend(ambiguous_term_comments(" ".join([title, action, expected_result, transcript])))
    if step.get("screenshotReviewStatus") not in (None, "", "approved"):
        reviewer_comments.append(f"Screenshot review status: {step.get('screenshotReviewStatus')}")
    if is_likely_non_application_screenshot(screenshot):
        reviewer_comments.append(
            "Selected screenshot appears to be a non-application Teams or meeting frame. Replace it with an application workflow screenshot before publishing."
        )
        screenshot = None
        caption = ""
    return GuideStep(
        title=title,
        action=action,
        expected_result=expected_result,
        section_title=text_value(step, "sectionTitle", "_sectionTitle", "section"),
        notes=notes,
        ui_text=ui_text,
        action_hints=action_hints,
        transcript=transcript,
        start=text_value(step, "start", "start_time"),
        end=text_value(step, "end", "end_time"),
        screenshot=screenshot,
        screenshot_caption=caption,
        reviewer_comments=reviewer_comments,
    )


def normalize_input(data: dict[str, Any], input_path: Path) -> GuideDraft:
    raw_data = data
    if isinstance(data.get("guideDraft"), dict):
        wrapped_model = data.get("model") if isinstance(data.get("model"), dict) else {}
        wrapped_generation_metadata = generation_metadata(data)
        data = data["guideDraft"]
        data.setdefault("model", {}).update(wrapped_model)
        data["_generationMetadata"] = {**generation_metadata(data), **wrapped_generation_metadata}

    document = data.get("document") if isinstance(data.get("document"), dict) else {}
    session = data.get("session") if isinstance(data.get("session"), dict) else {}
    recording = data.get("recording") if isinstance(data.get("recording"), dict) else {}
    source_recording = data.get("sourceRecording") if isinstance(data.get("sourceRecording"), dict) else {}
    metadata = data.get("metadata") if isinstance(data.get("metadata"), dict) else {}
    meta = data.get("meta") if isinstance(data.get("meta"), dict) else {}

    raw_steps = data.get("steps")
    if not isinstance(raw_steps, list):
        raw_steps = data.get("procedure_steps")
    if not isinstance(raw_steps, list):
        raw_steps = data.get("segments")
    if not isinstance(raw_steps, list):
        raw_steps = data.get("confirmedSteps")
    if not isinstance(raw_steps, list) or not raw_steps:
        raw_steps = data.get("pendingStepShells")
    if not isinstance(raw_steps, list) or not raw_steps:
        raw_steps = flatten_section_steps(data)
    if not isinstance(raw_steps, list):
        raise ValueError("Input JSON must contain a steps, procedure_steps, or segments array.")

    asset_roots = infer_asset_roots(data, input_path)
    title = text_value(
        document,
        "title",
        default=text_value(
            data,
            "title",
            default=text_value(
                data,
                "title",
                default=f"{text_value(recording, 'targetApplication', default=text_value(session, 'app_name', default='Application'))} User Guide",
            ),
        ),
    )
    version = text_value(document, "version", default=text_value(data, "version", default="Draft v0.1"))
    status = text_value(document, "status", default=text_value(data, "status", default="Prototype"))
    owner = text_value(document, "owner", default=text_value(data, "owner", default="KCXDocumentor"))
    effective_date = text_value(
        document,
        "effective_date",
        default=text_value(data, "effective_date", default=date.today().isoformat()),
    )
    application_name = first_text_value(
        (document, ("targetApplication", "application", "appName", "app_name")),
        (source_recording, ("targetApplication", "application", "appName", "app_name")),
        (recording, ("targetApplication", "application", "appName", "app_name")),
        (data, ("targetApplication", "application", "appName", "app_name")),
        (meta, ("targetApplication", "application", "appName", "app_name")),
        (session, ("targetApplication", "application", "appName", "app_name")),
        (metadata, ("targetApplication", "application", "appName", "app_name")),
        default="Not specified",
    )
    raw_duration = pick_value(document, "durationSeconds", "duration")
    if raw_duration is None:
        raw_duration = pick_value(source_recording, "durationSeconds", "duration")
    if raw_duration is None:
        raw_duration = pick_value(recording, "durationSeconds", "duration")
    if raw_duration is None:
        raw_duration = pick_value(session, "duration_sec", "duration")
    if raw_duration is None:
        raw_duration = pick_value(metadata, "durationSeconds", "duration")
    source_metadata = {
        "Application": application_name,
        "Recording Duration": format_duration(raw_duration),
        "Recorded At": text_value(session, "recorded_at", default=text_value(metadata, "recorded_at", default="Not specified")),
        "Source Recording": text_value(
            document,
            "sourceRecording",
            "sourceFile",
            default=text_value(source_recording, "sourceFile", default=text_value(recording, "sourceFile", default=text_value(meta, "sessionId", default=str(input_path)))),
        ),
        "Input File": str(input_path),
    }

    review_notes = (
        as_list(data.get("review_notes") or document.get("review_notes"))
        + as_list(data.get("warnings"))
        + as_list(data.get("assumptions"))
        + as_list(data.get("openReviewItems"))
        + as_list(data.get("reviewGuidance"))
        + as_list(data.get("reviewerGuidance"))
        + as_list(data.get("reviewComments"))
    )
    summary = customer_safe_summary(
        data,
        document,
        recording,
        source_recording,
        session,
        meta,
        metadata,
        title,
    )
    prerequisites = clean_prerequisites(as_list(data.get("prerequisites") or document.get("prerequisites")), application_name)

    return GuideDraft(
        title=title,
        version=version,
        status=text_value(meta, "draftStatus", default=status),
        owner=owner,
        effective_date=effective_date,
        summary=summary,
        audience=as_list(data.get("audience") or document.get("audience") or ["Application users"]),
        prerequisites=prerequisites,
        workflow_overview=clean_customer_items(
            as_list(
                data.get("workflow_overview")
                or data.get("overview")
                or document.get("workflow_overview")
                or section_body(data, "Workflow Overview")
            )
        ),
        expected_results=clean_customer_items(as_list(data.get("expected_results") or document.get("expected_results"))),
        troubleshooting=clean_customer_items(as_list(data.get("troubleshooting") or document.get("troubleshooting"))),
        review_notes=review_notes,
        source_metadata=source_metadata,
        generation_metadata=data.get("_generationMetadata") if isinstance(data.get("_generationMetadata"), dict) else generation_metadata(raw_data if raw_data is not data else data),
        steps=[normalize_step(step, idx + 1, input_path, asset_roots) for idx, step in enumerate(raw_steps) if isinstance(step, dict)],
    )


def infer_asset_roots(data: dict[str, Any], input_path: Path) -> list[Path]:
    roots: list[Path] = []
    for session_id in candidate_session_ids(data, input_path):
        processed_dir = WORKSPACE / "samples" / "processed" / session_id
        if processed_dir.exists():
            roots.append(processed_dir)
    source_recording = data.get("sourceRecording") if isinstance(data.get("sourceRecording"), dict) else {}
    source_file = text_value(source_recording, "sourceFile")
    if source_file:
        source_path = Path(source_file)
        if source_path.exists():
            roots.append(source_path.parent)
    return roots


def candidate_session_ids(data: dict[str, Any], input_path: Path) -> list[str]:
    ids = [
        text_value(data, "sessionId"),
        text_value(data.get("session") if isinstance(data.get("session"), dict) else {}, "sessionId", "id"),
        input_path.parent.name if input_path.parent.name else "",
    ]
    return list(dict.fromkeys(session_id for session_id in ids if session_id))


def flatten_section_steps(data: dict[str, Any]) -> list[dict[str, Any]]:
    sections = data.get("sections")
    if not isinstance(sections, list):
        return []
    flattened: list[dict[str, Any]] = []
    for section in sections:
        if not isinstance(section, dict):
            continue
        section_title = text_value(section, "title", default="Guide Section")
        steps = section.get("steps")
        if isinstance(steps, list) and steps:
            for step in steps:
                if not isinstance(step, dict):
                    continue
                enriched = dict(step)
                title = text_value(enriched, "title", "name")
                if not title:
                    instruction = text_value(enriched, "instruction", "action", "body", "summary")
                    enriched["title"] = derive_title_from_action(instruction, len(flattened) + 1) if instruction else section_title
                enriched.setdefault("_sectionTitle", section_title)
                flattened.append(enriched)
            continue
        body_items = as_list(section.get("body") or section.get("bullets") or section.get("summary"))
        for item in body_items:
            flattened.append({"title": derive_title_from_action(item, len(flattened) + 1), "instruction": item, "_sectionTitle": section_title})
    return flattened


def section_body(data: dict[str, Any], title: str) -> list[str]:
    sections = data.get("sections")
    if not isinstance(sections, list):
        return []
    for section in sections:
        if not isinstance(section, dict):
            continue
        if str(section.get("title", "")).strip().lower() == title.lower():
            return as_list(section.get("body")) + as_list(section.get("bullets"))
    return []


def add_labeled_paragraph(doc: Document, label: str, value: str) -> None:
    if not value:
        return
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.name = FONT
    add_rich_text(paragraph, value)


def add_optional_bullet_section(doc: Document, title: str, items: list[str]) -> None:
    if not items:
        return
    doc.add_heading(title, level=1)
    add_rich_bullets(doc, items)


def add_required_bullet_section(doc: Document, title: str, items: list[str], fallback: str) -> None:
    doc.add_heading(title, level=1)
    add_rich_bullets(doc, items or [fallback])


def add_reviewer_comment(doc: Document, paragraph: Any, comments: list[str]) -> bool:
    clean_comments = [comment.strip() for comment in comments if comment and comment.strip()]
    if not clean_comments or not paragraph.runs or not hasattr(doc, "add_comment"):
        return False
    try:
        doc.add_comment(
            paragraph.runs,
            text="\n".join(clean_comments),
            author="KCXDocumentor Reviewer",
            initials="KCX",
        )
    except Exception:
        return False
    return True


def add_fallback_reviewer_section(doc: Document, comments: list[str]) -> None:
    if not comments:
        return
    doc.add_heading("Reviewer Comments", level=1)
    add_rich_bullets(doc, comments)


def render_step(doc: Document, step: GuideStep, index: int) -> list[str]:
    heading = doc.add_heading(f"{index}. {step.title}", level=2)
    step_comments = list(dict.fromkeys(step.reviewer_comments))
    if step.start or step.end:
        step_comments.append(f"Source timing: {' - '.join(part for part in [step.start, step.end] if part)}")
    if not step.screenshot:
        step_comments.append(f"Step {index}: No screenshot was available for this procedure step.")
    if not add_reviewer_comment(doc, heading, step_comments):
        fallback_comments = [f"Step {index}: {comment}" for comment in step_comments]
    else:
        fallback_comments = []

    add_labeled_paragraph(doc, "Action", step.action)
    add_labeled_paragraph(doc, "Expected result", step.expected_result)

    if step.screenshot:
        picture_p = doc.add_paragraph()
        picture_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            picture_p.add_run().add_picture(str(step.screenshot), width=Inches(6.1))
        except Exception as exc:
            add_labeled_paragraph(doc, "Screenshot unavailable", f"{step.screenshot} ({exc})")
        else:
            caption = add_rich_paragraph(doc, step.screenshot_caption or f"Step {index} screenshot")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in caption.runs:
                run.italic = True
                run.font.name = FONT
                run.font.size = Pt(8.8)
    else:
        pass

    return fallback_comments


def render_docx(draft: GuideDraft, output_path: Path) -> None:
    doc = Document()
    style_document(doc)
    add_header_footer(doc, draft.title, "keycentrix confidential")
    add_cover(doc, draft.title, draft.version, draft.status, draft.summary)
    add_metadata_table(
        doc,
        [
            ("Owner", draft.owner),
            ("Effective Date", draft.effective_date),
            ("Status", draft.status),
            ("Version", draft.version),
        ],
    )

    doc.add_heading("Purpose", level=1)
    purpose_p = add_rich_paragraph(doc, draft.summary)
    fallback_comments: list[str] = []
    if draft.review_notes and not add_reviewer_comment(doc, purpose_p, draft.review_notes):
        fallback_comments.extend(draft.review_notes)

    add_optional_bullet_section(doc, "Intended Audience", draft.audience)
    add_required_bullet_section(
        doc,
        "Prerequisites",
        draft.prerequisites,
        "Confirm access to the target application and review the source recording context before publishing.",
    )
    add_required_bullet_section(
        doc,
        "Workflow Overview",
        draft.workflow_overview,
        "Follow the workflow sections in order and confirm the expected screen or request status before moving to the next step.",
    )

    doc.add_heading("Step-by-Step Procedures", level=1)
    active_section = ""
    for index, step in enumerate(draft.steps, start=1):
        if step.section_title and step.section_title != active_section:
            doc.add_heading(step.section_title, level=2)
            active_section = step.section_title
        fallback_comments.extend(render_step(doc, step, index))

    add_required_bullet_section(
        doc,
        "Expected Results",
        draft.expected_results,
        "The user can complete the documented workflow and confirm the expected request or configuration outcome.",
    )
    add_required_bullet_section(
        doc,
        "Troubleshooting Notes",
        draft.troubleshooting,
        "If the workflow does not match the current application behavior, confirm the current process with the appropriate application owner.",
    )
    add_fallback_reviewer_section(doc, fallback_comments)

    doc.add_heading("Appendix: Source Recording Metadata", level=1)
    add_metadata_table(doc, list(draft.source_metadata.items()))

    apply_markdown_bold(doc)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Render a keycentrix-styled DOCX guide from guide draft JSON or procedure_trace JSON."
    )
    parser.add_argument("input", nargs="?", type=Path, help="Path to guide draft JSON or procedure_trace JSON.")
    parser.add_argument(
        "--input",
        dest="input_option",
        type=Path,
        help="Path to guide draft JSON or procedure_trace JSON. Equivalent to the positional input.",
    )
    parser.add_argument(
        "-o",
        "--output",
        type=Path,
        default=WORKSPACE / "artifacts" / "generated" / "prototype-guide.docx",
        help="Output DOCX path. Defaults to artifacts/generated/prototype-guide.docx.",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    raw_input = args.input_option or args.input
    if raw_input is None:
        print("Input file is required. Pass a positional input or --input.", file=sys.stderr)
        return 2
    input_path = raw_input.resolve()
    if not input_path.exists():
        print(f"Input file not found: {input_path}", file=sys.stderr)
        return 2

    try:
        data = json.loads(input_path.read_text(encoding="utf-8"))
        draft = normalize_input(data, input_path)
        render_docx(draft, args.output.resolve())
    except Exception as exc:
        print(f"Failed to render DOCX: {exc}", file=sys.stderr)
        return 1

    print(f"Wrote {args.output.resolve()}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
