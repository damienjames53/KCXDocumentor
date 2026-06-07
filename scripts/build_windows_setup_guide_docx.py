#!/usr/bin/env python3
from __future__ import annotations

import argparse
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from build_user_guide_docx import (
    INK,
    LOGO_CANDIDATES,
    MUTED,
    PRIMARY_BLUE,
    PRIMARY_STRONG,
    WHITE,
    add_header_footer,
    first_existing,
    markdown_to_doc,
    set_cell_margins,
    set_cell_shading,
    set_document_styles,
    set_table_width,
    shade_paragraph,
)


WORKSPACE = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = WORKSPACE / "docs" / "windows-setup-guide.md"
DEFAULT_OUTPUT = WORKSPACE / "docs" / "KCXDocumentor Setup Guide For Windows.docx"


def add_cover(doc: Document) -> None:
    logo = first_existing(LOGO_CANDIDATES)
    if logo:
        p = doc.add_paragraph()
        p.add_run().add_picture(str(logo), width=Inches(2.35))

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    run = p.add_run("KCXDocumentor Setup Guide For Windows")
    run.font.name = "Aptos Display"
    run.font.size = Pt(25)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_STRONG

    p = doc.add_paragraph()
    run = p.add_run("Docker Desktop, private GHCR image access, local folder mapping, first-run validation, and initial guide generation.")
    run.font.name = "Aptos"
    run.font.size = Pt(11.5)
    run.font.italic = True
    run.font.color.rgb = MUTED

    p = doc.add_paragraph()
    shade_paragraph(p, "1C75BC")
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Audience: internal testers, business analysts, trainers, implementation teams, and technical setup helpers.")
    run.font.name = "Aptos"
    run.font.size = Pt(10.5)
    run.font.bold = True
    run.font.color.rgb = WHITE

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, [2100, 6000])
    rows = [
        ("Document type", "Windows setup and first-run guide"),
        ("Application", "KCXDocumentor"),
        ("Container image", "ghcr.io/keycentrix/kcxdocumentor:latest"),
        ("Local URL", "http://127.0.0.1:8765"),
        ("Revision date", date.today().isoformat()),
    ]
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], "EAF3FB")
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.color.rgb = INK
        row.cells[1].paragraphs[0].runs[0].font.color.rgb = INK

    doc.add_page_break()


def keep_table_rows_together(doc: Document) -> None:
    for table in doc.tables:
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            if tr_pr.find(qn("w:cantSplit")) is None:
                tr_pr.append(OxmlElement("w:cantSplit"))


def build(source: Path, output: Path) -> None:
    doc = Document()
    set_document_styles(doc)
    add_header_footer(doc, "KCXDocumentor Windows Setup")
    add_cover(doc)
    markdown_to_doc(doc, source)
    keep_table_rows_together(doc)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser(description="Build the KCXDocumentor Windows setup guide DOCX.")
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()

    build(args.source.resolve(), args.output.resolve())
    print(f"Built {args.output.resolve()}")


if __name__ == "__main__":
    main()
