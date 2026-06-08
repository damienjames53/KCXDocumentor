#!/usr/bin/env python3
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document

from build_prototype_collateral_docx import (
    DocSpec,
    OUT_DIR,
    _bullets,
    _callout,
    _cover,
    _header_footer,
    _style_document,
    _table,
)


# Keep the original SharePoint filename for link continuity. The document title
# has been updated to reflect the current Azure Foundry provider path.
OUTPUT = OUT_DIR / "KCXDocumentor - Anthropic API Data Protection Brief.docx"


def add_doc_control(doc: Document, spec: DocSpec) -> None:
    doc.add_heading("Document Control", level=1)
    _table(
        doc,
        ["Control", "Value"],
        [
            ["Document Title", spec.title],
            ["Document ID", spec.document_id],
            ["Status", spec.status],
            ["Owner", spec.owner],
            ["Audience", spec.audience],
            ["Review Basis", "Microsoft Foundry Claude and Anthropic API documentation reviewed on 2026-06-07"],
        ],
    )
    doc.add_heading("Version History", level=1)
    _table(
        doc,
        ["Version", "Date", "Summary"],
        [
            ["1.0", "2026-06-05", "Initial first-party Anthropic API data protection and BAA readiness brief."],
            ["2.0", date.today().isoformat(), "Updated for Azure Foundry Claude, Azure Function proxy, Cosmos usage reporting, and Marketplace billing caveats."],
        ],
    )


def build() -> Path:
    spec = DocSpec(
        title="KCXDocumentor - Azure Foundry Claude Data Protection Brief",
        subtitle="Executive review of AI data flow, PHI/PII exposure, Marketplace billing, and compliance readiness",
        document_id="KCXDOC-DATA-001",
        status="Executive review",
        owner="keycentrix Engineering and Product",
        audience="Executive Team, Security, Compliance, Product, Training, Implementation, and Engineering",
        purpose="Summarize the current KCXDocumentor AI provider path and the protections, caveats, and product controls needed before the application is used with protected health information or sensitive personal information.",
        scope="Focused on KCXDocumentor guide generation through Microsoft Foundry Claude Sonnet 4.6 behind an authenticated Azure Function proxy. This brief does not evaluate Claude consumer products, Claude Team or Enterprise product interfaces, direct Anthropic API production use, or unrelated third-party integrations.",
        sections=[],
        output_name=OUTPUT.name,
    )

    doc = Document()
    doc.core_properties.title = spec.title
    doc.core_properties.author = "keycentrix"
    doc.core_properties.subject = spec.subtitle
    doc.core_properties.comments = "Executive compliance brief built with the DamienDev document recipe pattern."
    _style_document(doc)
    _header_footer(doc, spec.title)
    _cover(doc, spec)
    add_doc_control(doc, spec)

    doc.add_heading("Purpose", level=1)
    doc.add_paragraph(spec.purpose)
    doc.add_heading("Scope", level=1)
    doc.add_paragraph(spec.scope)
    doc.add_heading("Intended Audience", level=1)
    doc.add_paragraph(spec.audience)

    doc.add_heading("Executive Summary", level=1)
    _callout(
        doc,
        "Recommended executive position",
        "KCXDocumentor may continue the pilot with synthetic, internal, or de-identified recordings using the current Azure Foundry Claude path. PHI-bearing production use should wait until Security and Compliance confirm the selected Foundry Claude Marketplace terms, Anthropic processor terms, Microsoft DPA posture, region behavior, and PHI-mode controls are acceptable.",
    )
    _bullets(
        doc,
        [
            "The desktop app no longer calls Anthropic directly. It sends compact prompt data to the authenticated Azure Function configured for Azure Foundry Claude.",
            "Raw recordings, extracted audio, frame sets, OCR artifacts, processed traces, DOCX files, and local QA artifacts remain on the workstation.",
            "The Azure Function validates the signed-in user, schedules generation, calls Azure Foundry Claude, records success or failure usage, and stores AI Spend metadata in Cosmos DB.",
            "Microsoft documentation for Claude in Foundry states that Anthropic is the processor for prompts and outputs, while Microsoft manages the API deployment infrastructure and billing/usage processing.",
            "Claude in Foundry is a Microsoft Marketplace model path. Marketplace billing and model-publisher terms must be reviewed separately from standard Azure credits or ordinary Azure service assumptions.",
            "For PHI/PII risk, the sensitive transmission points are transcript excerpts, OCR text, reviewer notes, approved screenshot context, and generated guide content, not the raw video file.",
        ],
    )

    doc.add_heading("Current KCXDocumentor AI Path", level=1)
    _table(
        doc,
        ["Layer", "Current Handling", "Protection / Caveat"],
        [
            ["Desktop app", "Runs locally in Docker on the user's workstation.", "No provider key is stored on the workstation. The app authenticates the user with Microsoft Entra MSAL + PKCE."],
            ["Local processing", "FFmpeg, Whisper, OCR, frame scoring, frame review, DOCX build, and QA run locally.", "Large media and guide artifacts stay in mapped local folders."],
            ["Prompt payload", "Only compact reviewed context is sent for generation.", "Prompt data can still contain PHI/PII if the recording, OCR, transcript, or reviewer notes contain it."],
            ["Azure Function", "Acts as the server-side policy boundary.", "Validates user tokens, queues generation, enforces token scheduling, calls Foundry, and records usage/failure details."],
            ["Azure Foundry Claude", "Uses Claude Sonnet 4.6 through the Foundry Marketplace deployment.", "Anthropic processes prompts and outputs for the Claude API; Microsoft manages deployment infrastructure and usage/billing processing."],
            ["Cosmos DB AI Spend", "Stores usage, document, owner, page count, cost estimate, status, and failure metadata.", "Usage records must avoid PHI in titles, session names, failure reasons, or free-text metadata."],
        ],
    )

    doc.add_heading("Data Flow Risk Considerations", level=1)
    _table(
        doc,
        ["Data Element", "Current Handling", "Executive Risk View"],
        [
            ["Raw video recording", "Processed locally on the workstation.", "Lowest vendor exposure if raw video is never sent to the Azure Function or model provider."],
            ["Speech transcript", "Generated locally with Whisper or imported as a sidecar and summarized into prompt context.", "Can contain PHI/PII if narration includes patient, prescription, account, or support details."],
            ["OCR and screen text", "Extracted locally and used as visible UI evidence.", "Can expose patient, pharmacy, employee, account, or workflow-sensitive details shown on screen."],
            ["Reviewer notes", "Used to steer guide generation and explain uncertainty.", "Can introduce PHI/PII if reviewers type identifiers, case details, or customer-specific context."],
            ["Approved screenshots", "Selected locally for guide context and final DOCX output.", "Image content can include PHI/PII if screens are not masked, de-identified, or reviewed."],
            ["Prompt payload", "Sent to the Azure Function and then to Azure Foundry Claude.", "Primary cloud data exposure point; must be minimized and governed."],
            ["Generated guide output", "Returned by Claude, built into a local DOCX, and optionally downloaded by the user.", "May restate sensitive content unless prompts, QA, and human review prevent it."],
            ["Usage reporting", "Stored in Cosmos DB for AI Spend visibility.", "Should remain operational metadata only and should not include patient or customer identifiers."],
        ],
    )

    doc.add_heading("Published Protections And Caveats", level=1)
    _table(
        doc,
        ["Topic", "Published Position", "Meaning For KCXDocumentor"],
        [
            ["Azure Foundry Claude data processing", "Microsoft states that for Claude in Foundry, Anthropic acts as the data processor for prompts and outputs, while Microsoft manages the API deployment infrastructure.", "Azure routing does not eliminate Anthropic data processing. Compliance review must consider both Microsoft and Anthropic terms."],
            ["Regional processing", "Microsoft states that Claude prompts and outputs may be processed outside the selected region for operational purposes.", "Do not assume single-region processing for PHI-sensitive workflows without formal compliance approval."],
            ["Marketplace transaction", "Microsoft states that Marketplace contact, transaction, billing, and usage details may be shared with the model publisher.", "Marketplace billing and publisher visibility must be reviewed before moving to a production subscription."],
            ["Marketplace terms", "Microsoft Product Terms state that Azure Marketplace products are subject to separate Marketplace terms.", "Standard Azure credits or consumption commitments may not cover Claude Marketplace usage."],
            ["Anthropic retention and training", "Anthropic documents API data-retention options, ZDR arrangements, and that retained API data is not used for model training without permission.", "This supports confidential use, but does not by itself approve PHI workflows."],
            ["Anthropic HIPAA-ready API", "Anthropic documents HIPAA-ready API access with a signed BAA and a HIPAA-enabled organization for eligible Claude API features.", "This is relevant but must be mapped carefully to the Foundry Marketplace path before PHI use is approved."],
            ["Unsupported features", "Anthropic documents that some API features are not ZDR or HIPAA eligible.", "KCXDocumentor should keep guide generation limited to Messages-style behavior and avoid Files API, tools, web fetch, MCP connectors, and agent features for PHI workflows unless explicitly approved."],
        ],
    )

    doc.add_heading("Required Controls Before PHI Use", level=1)
    _bullets(
        doc,
        [
            "Create a PHI mode that blocks guide generation unless the approved provider path, terms, and compliance controls are active.",
            "Add a preflight warning that identifies transcript, OCR, reviewer notes, approved screenshots, document titles, and usage metadata as possible PHI/PII exposure points.",
            "Add local redaction or masking options for patient names, dates of birth, prescription identifiers, phone numbers, addresses, email addresses, MRNs, account numbers, and customer-specific support identifiers before prompt assembly.",
            "Keep API keys only in Azure Function App settings. Do not store Anthropic or Azure Foundry provider keys on workstations.",
            "Keep raw recordings, extracted frames, processed traces, generated DOCX files, and QA output in mapped local folders with workstation access controls.",
            "Log only operational metadata needed for audit and cost reporting. Avoid PHI in cloud usage records, file names, session names, generated titles, or telemetry.",
            "Disable or block unsupported model features for PHI workflows, including Files API, batch processing, web fetch, external tools, MCP connectors, code execution, and non-covered beta features unless Compliance explicitly approves them.",
            "Document a user-facing pilot rule: recordings used for guide generation must be synthetic, internal-only, or de-identified unless the PHI-ready path is approved.",
        ],
    )

    doc.add_heading("Decision Matrix", level=1)
    _table(
        doc,
        ["Operating Scenario", "Allowed For Pilot?", "Allowed For PHI?", "Recommended Control"],
        [
            ["Internal demo with synthetic or de-identified data", "Yes", "No PHI present", "Use current Azure Function and Azure Foundry Claude path with local-first controls and reviewer QA."],
            ["Internal workflow recording with possible PII but no PHI", "Conditional", "Not applicable", "Minimize, redact, and confirm business approval before generation."],
            ["Recording that may show PHI", "No, unless de-identified", "Not yet approved", "Require Compliance approval of Foundry Claude terms, PHI mode, redaction checks, and logging controls."],
            ["Production customer-facing documentation from real pharmacy workflows", "Conditional", "Only with covered controls", "Use the approved PHI-ready provider path or ensure complete de-identification before generation."],
            ["Direct first-party Anthropic API production path", "No", "No", "Do not use unless a separate approved architecture, BAA, and key-management model are established."],
        ],
    )

    doc.add_heading("Executive Decision Points", level=1)
    _bullets(
        doc,
        [
            "Confirm whether KCXDocumentor pilot recordings must be fully de-identified or may include limited non-PHI business context.",
            "Assign Security and Compliance review of the Azure Foundry Claude Marketplace terms, Anthropic processor terms, Microsoft DPA posture, region behavior, and Marketplace billing implications.",
            "Decide whether the organization requires a formal Anthropic BAA path, a Microsoft/Marketplace confirmation path, or both before PHI-bearing production use.",
            "Approve or defer PHI mode, local redaction checks, metadata hygiene checks, and unsupported-feature blocking.",
            "Confirm that the target Azure subscription payment method covers Marketplace charges before moving production usage to a company subscription.",
        ],
    )

    doc.add_heading("Source References", level=1)
    _table(
        doc,
        ["Source", "Reviewed Topic"],
        [
            ["Microsoft Learn: Data, privacy, and security for Anthropic Claude models in Microsoft Foundry, updated 2026-05-18", "Anthropic processor role, Microsoft deployment infrastructure role, Marketplace information sharing, and regional processing caveat."],
            ["Microsoft Product Terms for Microsoft Azure", "Marketplace terms, Foundry Models distinctions, and separate terms for third-party models and Marketplace products."],
            ["Anthropic API and Data Retention documentation", "Standard retention, zero data retention, HIPAA-ready API access, PHI handling guidance, and unsupported feature eligibility."],
            ["Anthropic Commercial Terms and Data Processing Addendum", "Customer content ownership, model-training restrictions, and Anthropic processor terms for commercial API use."],
            ["KCXDocumentor implementation plan", "Current local-first architecture, Azure Function proxy, Azure Foundry provider settings, Cosmos usage reporting, and queue/token controls."],
        ],
    )

    doc.add_heading("Recipe Alignment", level=1)
    doc.add_paragraph(
        "Built as a keycentrix internal artifact using the DamienDev document recipe structure: control metadata, version history, purpose, scope, audience, governance, process detail, decision matrix, and source appendix."
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


def main() -> int:
    output = build()
    print(f"Built {output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
