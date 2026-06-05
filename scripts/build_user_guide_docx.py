#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


WORKSPACE = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = WORKSPACE / "docs" / "user-guide.md"
DEFAULT_OUTPUT = WORKSPACE / "docs" / "user-guide.docx"
LOGO_CANDIDATES = [
    WORKSPACE / "assets" / "branding" / "images" / "keycentrix-template-logo.png",
    WORKSPACE / "assets" / "branding" / "images" / "keycentrix-full-logo.png",
    WORKSPACE / "assets" / "branding" / "images" / "sticky-logo.png",
]

PRIMARY_BLUE = RGBColor(0x1C, 0x75, 0xBC)
PRIMARY_STRONG = RGBColor(0x12, 0x3D, 0x75)
GREEN = RGBColor(0x8C, 0xC6, 0x3F)
INK = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

INLINE_RE = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\[[^\]]+\]\(([^)]+)\))")


def first_existing(paths: list[Path]) -> Path | None:
    return next((path for path in paths if path.exists()), None)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, margin: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def create_numbering_sequence(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
        if element.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(element.get(qn("w:numId")))
        for element in numbering.findall(qn("w:num"))
        if element.get(qn("w:numId")) is not None
    ]
    abstract_id = (max(abstract_ids) + 1) if abstract_ids else 1
    num_id = (max(num_ids) + 1) if num_ids else 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "decimal")
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "%1.")
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    lvl.extend([start, fmt, text, jc, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def set_document_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in [
        ("Heading 1", 18, PRIMARY_STRONG),
        ("Heading 2", 14, PRIMARY_BLUE),
        ("Heading 3", 11.5, PRIMARY_STRONG),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Aptos Display"
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(5)

    for style_name in ["List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(10)


def add_header_footer(doc: Document, header_title: str) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(header_title)
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("keycentrix confidential and intended for the audience named in this guide.")
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED


def add_inline(paragraph, text: str) -> None:
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        elif token.startswith("["):
            label = token[1 : token.index("](")]
            run = paragraph.add_run(label)
            run.font.color.rgb = PRIMARY_BLUE
            run.underline = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_cover(doc: Document, title: str, subtitle: str, audience: str) -> None:
    logo = first_existing(LOGO_CANDIDATES)
    if logo:
        p = doc.add_paragraph()
        p.add_run().add_picture(str(logo), width=Inches(2.35))

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    run = p.add_run(title)
    run.font.name = "Aptos Display"
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_STRONG

    p = doc.add_paragraph()
    run = p.add_run(subtitle)
    run.font.name = "Aptos"
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = MUTED

    p = doc.add_paragraph()
    shade_paragraph(p, "1C75BC")
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(audience)
    run.font.name = "Aptos"
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = WHITE

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, [2100, 6000])
    rows = [
        ("Document type", "Application user guide"),
        ("Application", "KCXDocumentor"),
        ("Revision date", date.today().isoformat()),
        ("Source", "docs/user-guide.md and current local UI screenshots"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], "EAF3FB")
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
        row.cells[0].paragraphs[0].runs[0].font.bold = True

    doc.add_page_break()


def add_screenshot(doc: Document, image_path: Path, caption: str) -> None:
    if not image_path.exists():
        p = doc.add_paragraph()
        run = p.add_run(f"Screenshot unavailable: {caption}")
        run.italic = True
        run.font.color.rgb = MUTED
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(image_path), width=Inches(6.35))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(caption)
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [int(8100 / len(rows[0])) for _ in rows[0]]
    set_table_width(table, widths)
    for idx, header in enumerate(rows[0]):
        table.rows[0].cells[idx].text = header
        set_cell_shading(table.rows[0].cells[idx], "1C75BC")
        run = table.rows[0].cells[idx].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = WHITE
    for data in rows[1:]:
        cells = table.add_row().cells
        for idx, value in enumerate(data):
            cells[idx].text = value
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def markdown_to_doc(doc: Document, source: Path) -> None:
    table_buffer: list[list[str]] = []
    current_num_id: int | None = None
    skip_title = True

    def flush_table() -> None:
        nonlocal table_buffer
        if table_buffer:
            add_table(doc, table_buffer)
            table_buffer = []

    def reset_numbering() -> None:
        nonlocal current_num_id
        current_num_id = None

    lines = source.read_text(encoding="utf-8").splitlines()
    in_code = False
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("```"):
            flush_table()
            reset_numbering()
            in_code = not in_code
            continue
        if in_code:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(8.8)
            continue
        if not stripped:
            flush_table()
            reset_numbering()
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            reset_numbering()
            cells = [cell.strip() for cell in stripped.strip("|").split("|")]
            if all(set(cell) <= {"-", ":"} for cell in cells):
                continue
            table_buffer.append(cells)
            continue

        flush_table()

        if stripped.startswith("![") and "](" in stripped and stripped.endswith(")"):
            reset_numbering()
            caption = stripped[2 : stripped.index("](")]
            target = stripped[stripped.index("](") + 2 : -1]
            add_screenshot(doc, (source.parent / target).resolve(), caption)
        elif stripped.startswith("# "):
            reset_numbering()
            if skip_title:
                skip_title = False
                continue
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, stripped[2:])
        elif stripped.startswith("## "):
            reset_numbering()
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, stripped[3:])
        elif stripped.startswith("### "):
            reset_numbering()
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, stripped[4:])
        elif stripped.startswith("#### "):
            reset_numbering()
            p = doc.add_paragraph(style="Heading 3")
            add_inline(p, stripped[5:])
        elif stripped.startswith("- "):
            reset_numbering()
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, stripped[2:])
        elif re.match(r"^\\d+\\. ", stripped):
            if current_num_id is None:
                current_num_id = create_numbering_sequence(doc)
            p = doc.add_paragraph()
            apply_numbering(p, current_num_id)
            add_inline(p, re.sub(r"^\\d+\\. ", "", stripped))
        else:
            reset_numbering()
            p = doc.add_paragraph()
            add_inline(p, stripped)

    flush_table()


def build(source: Path, output: Path) -> None:
    doc = Document()
    set_document_styles(doc)
    add_header_footer(doc, "KCXDocumentor User Guide")
    add_cover(
        doc,
        "KCXDocumentor User Guide",
        "Local recording review, guide generation, DOCX download, QA review, and AI Spend tracking.",
        "Audience: business analysts, trainers, implementation teams, and documentation reviewers.",
    )
    markdown_to_doc(doc, source)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser(description="Build the KCXDocumentor user guide DOCX.")
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()

    build(args.source.resolve(), args.output.resolve())
    print(f"Built {args.output.resolve()}")


if __name__ == "__main__":
    main()
