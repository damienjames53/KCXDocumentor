from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


WORKSPACE = Path(__file__).resolve().parents[2]
SHARED_IMAGES = WORKSPACE / "assets" / "branding" / "images"
LOGO = SHARED_IMAGES / "keycentrix-logo.png"

FONT = "Calibri"
INK = RGBColor(0x1F, 0x29, 0x37)
SLATE = RGBColor(0x6B, 0x72, 0x80)
BLUE = RGBColor(0x1C, 0x75, 0xBC)
DARK = RGBColor(0x1F, 0x29, 0x37)


def set_cell_width(cell, width_inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.1

    h1 = doc.styles["Heading 1"]
    h1.font.name = FONT
    h1.font.size = Pt(16)
    h1.font.color.rgb = BLUE
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(5)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = FONT
    h2.font.size = Pt(12.5)
    h2.font.color.rgb = DARK
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(3)
    h2.paragraph_format.keep_with_next = True


def add_page_number(paragraph) -> None:
    run = paragraph.add_run("Page ")
    run.font.name = FONT
    run.font.size = Pt(8.5)
    run.font.color.rgb = SLATE

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run_el = OxmlElement("w:r")
    run_el.append(OxmlElement("w:rPr"))
    run_el.append(fld_begin)
    run_el.append(instr)
    run_el.append(fld_end)
    paragraph._p.append(run_el)


def add_header_footer(doc: Document, title: str, footer_text: str) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(title)
    r.font.name = FONT
    r.font.size = Pt(8.5)
    r.font.color.rgb = SLATE

    footer = section.footer
    left = footer.paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = left.add_run(footer_text)
    r.font.name = FONT
    r.font.size = Pt(8.5)
    r.font.color.rgb = SLATE

    right = footer.add_paragraph()
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(right)


def add_cover(doc: Document, title: str, version: str, status: str, summary: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(24)
    if LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Inches(2.6))

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(18)
    r = title_p.add_run(title)
    r.font.name = FONT
    r.font.size = Pt(24)
    r.bold = True
    r.font.color.rgb = DARK

    version_p = doc.add_paragraph()
    r = version_p.add_run(version)
    r.font.name = FONT
    r.font.size = Pt(12)
    r.font.color.rgb = INK

    status_p = doc.add_paragraph()
    r = status_p.add_run(f"keycentrix - {status}")
    r.font.name = FONT
    r.font.size = Pt(11)
    r.bold = True
    r.font.color.rgb = INK

    summary_p = doc.add_paragraph()
    r = summary_p.add_run(summary)
    r.font.name = FONT
    r.font.size = Pt(10.5)
    r.font.color.rgb = INK


def add_metadata_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    doc.add_heading("Document Control Metadata", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    widths = [2.1, 4.65]
    table.rows[0].cells[0].text = "Metadata Field"
    table.rows[0].cells[1].text = "Document Value"
    for field, value in rows:
        cells = table.add_row().cells
        cells[0].text = field
        cells[1].text = value
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = FONT
                    run.font.size = Pt(9.8)
                    run.font.color.rgb = INK
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_section(doc: Document, number: int, title: str, paragraphs: list[str] | None = None, bullets: list[str] | None = None) -> None:
    doc.add_heading(f"{number}. {title}", level=1)
    for paragraph in paragraphs or []:
        doc.add_paragraph(paragraph)
    if bullets:
        add_bullets(doc, bullets)


def add_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cells[idx].text = value
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(1)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = FONT
                    run.font.size = Pt(8.8)
                    run.font.color.rgb = INK
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = DARK
